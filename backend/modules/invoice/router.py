"""
Invoice module endpoints: clienti, facturi, generare PDF, integrare AI,
rapoarte, export, template documente, email, scanner OCR.

Endpoints:
  GET    /api/invoice/clients          — Lista clienti
  POST   /api/invoice/clients          — Creare client
  GET    /api/invoice/clients/:id      — Detalii client
  PUT    /api/invoice/clients/:id      — Actualizare client
  DELETE /api/invoice/clients/:id      — Stergere client
  GET    /api/invoice/clients/:id/history — Istoric comenzi per client
  GET    /api/invoice/list             — Lista facturi (cu nume client)
  POST   /api/invoice/create           — Creare factura
  GET    /api/invoice/:id              — Detalii factura + client
  PUT    /api/invoice/:id              — Actualizare factura (doar draft)
  DELETE /api/invoice/:id              — Stergere factura (doar draft)
  POST   /api/invoice/:id/pdf         — Generare PDF factura
  PUT    /api/invoice/:id/status      — Schimbare status factura
  POST   /api/invoice/:id/send-email  — Trimite factura pe email
  POST   /api/invoice/generate-from-calc — Pre-fill factura din calcul AI
  GET    /api/invoice/reports/monthly  — Raport lunar venituri
  GET    /api/invoice/reports/summary  — Sumar total venituri
  GET    /api/invoice/export/csv       — Export facturi CSV
  GET    /api/invoice/export/excel     — Export facturi Excel
  GET    /api/invoice/templates        — Lista template documente
  POST   /api/invoice/templates/generate — Genereaza document din template
  POST   /api/invoice/scan             — Scanner facturi primite (OCR+AI)
"""

from __future__ import annotations

import csv
import io
import json
import logging
import os
import smtplib
import tempfile
from datetime import date, datetime, timezone
from email.mime.application import MIMEApplication

import httpx
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from pathlib import Path
from typing import Any, Optional

from fastapi import APIRouter, File, HTTPException, UploadFile
from fastapi.responses import FileResponse, StreamingResponse
from pydantic import BaseModel, Field

from app.config import settings
from app.core.activity_log import log_activity
from app.db.database import get_db

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/invoice", tags=["Facturare"])

# --- Directory for saved PDFs ---
INVOICES_DIR = settings.data_dir / "invoices"
INVOICES_DIR.mkdir(parents=True, exist_ok=True)


# ═══════════════════════════════════════════
# Pydantic models
# ═══════════════════════════════════════════

class ClientCreate(BaseModel):
    name: str
    cui: Optional[str] = None
    reg_com: Optional[str] = None
    address: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    notes: Optional[str] = None


class ClientUpdate(BaseModel):
    name: Optional[str] = None
    cui: Optional[str] = None
    reg_com: Optional[str] = None
    address: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    notes: Optional[str] = None


class InvoiceItem(BaseModel):
    description: str
    quantity: float = 1.0
    unit_price: float = 0.0
    total: float = 0.0


class InvoiceCreate(BaseModel):
    client_id: int
    items: list[InvoiceItem]
    series_prefix: Optional[str] = None
    date: str = Field(default_factory=lambda: date.today().isoformat())
    due_date: Optional[str] = None
    notes: Optional[str] = None
    vat_percent: float = 0.0


class OfferPdfRequest(BaseModel):
    client_name: str
    client_address: Optional[str] = None
    items: list[InvoiceItem]
    notes: Optional[str] = None
    validity_days: int = 30


class InvoiceUpdate(BaseModel):
    client_id: Optional[int] = None
    items: Optional[list[InvoiceItem]] = None
    date: Optional[str] = None
    due_date: Optional[str] = None
    notes: Optional[str] = None
    vat_percent: Optional[float] = None


class StatusUpdate(BaseModel):
    status: str  # draft, sent, paid, cancelled


class GenerateFromCalcRequest(BaseModel):
    calculation_id: int


class InvoiceSeriesCreate(BaseModel):
    prefix: str
    name: str
    description: str | None = None


class InvoiceSeriesUpdate(BaseModel):
    name: str | None = None
    description: str | None = None
    active: bool | None = None


class SendEmailRequest(BaseModel):
    to_email: str
    subject: Optional[str] = None
    body: Optional[str] = None


class TemplateGenerateRequest(BaseModel):
    template_name: str
    data: dict[str, Any]
    output_format: str = "docx"  # docx or pdf


# ═══════════════════════════════════════════
# Helper functions
# ═══════════════════════════════════════════

async def _next_invoice_number(series_prefix: str | None = None) -> str:
    """Generate next sequential invoice number using configurable series."""
    year = date.today().year
    async with get_db() as db:
        if not series_prefix:
            # Get default series prefix
            cursor = await db.execute(
                "SELECT prefix FROM invoice_series WHERE is_default = 1 AND active = 1 LIMIT 1"
            )
            row = await cursor.fetchone()
            series_prefix = row["prefix"] if row else "RCC"
        prefix = f"{series_prefix}-{year}-"
        cursor = await db.execute(
            "SELECT invoice_number FROM invoices WHERE invoice_number LIKE ? ORDER BY invoice_number DESC LIMIT 1",
            (f"{prefix}%",),
        )
        row = await cursor.fetchone()
    if row:
        try:
            last_num = int(row["invoice_number"].split("-")[-1])
        except (ValueError, IndexError):
            last_num = 0
        return f"{prefix}{last_num + 1:03d}"
    return f"{prefix}001"


def _calculate_totals(items: list[InvoiceItem], vat_percent: float) -> tuple[float, float, float]:
    """Calculate subtotal, VAT amount, and total from items."""
    subtotal = sum(item.quantity * item.unit_price for item in items)
    vat_amount = subtotal * vat_percent / 100.0
    total = subtotal + vat_amount
    return round(subtotal, 2), round(vat_amount, 2), round(total, 2)


def _items_to_json(items: list[InvoiceItem]) -> str:
    """Serialize items list to JSON string."""
    return json.dumps(
        [item.model_dump() for item in items],
        ensure_ascii=False,
    )


def _items_from_json(json_str: str) -> list[dict]:
    """Deserialize items JSON string."""
    try:
        return json.loads(json_str)
    except (json.JSONDecodeError, TypeError):
        return []


def _row_to_dict(row) -> dict:
    """Convert aiosqlite Row to dict, parsing items_json."""
    d = dict(row)
    if "items_json" in d:
        d["items"] = _items_from_json(d.pop("items_json"))
    return d


# ═══════════════════════════════════════════
# Client endpoints
# ═══════════════════════════════════════════

@router.get("/clients")
async def list_clients():
    """Lista toti clientii."""
    async with get_db() as db:
        cursor = await db.execute(
            "SELECT * FROM clients ORDER BY name ASC"
        )
        rows = await cursor.fetchall()
    return [dict(row) for row in rows]


@router.post("/clients", status_code=201)
async def create_client(data: ClientCreate):
    """Creare client nou."""
    async with get_db() as db:
        cursor = await db.execute(
            """INSERT INTO clients (name, cui, reg_com, address, email, phone, notes)
               VALUES (?, ?, ?, ?, ?, ?, ?)""",
            (data.name, data.cui, data.reg_com, data.address, data.email, data.phone, data.notes),
        )
        await db.commit()
        client_id = cursor.lastrowid

    await log_activity(
        action="invoice.client_create",
        summary=f"Client creat: {data.name}",
        details={"client_id": client_id, "name": data.name},
    )
    return {"id": client_id, "message": f"Client '{data.name}' creat cu succes."}


@router.get("/clients/{client_id}")
async def get_client(client_id: int):
    """Detalii client dupa ID."""
    async with get_db() as db:
        cursor = await db.execute(
            "SELECT * FROM clients WHERE id = ?", (client_id,)
        )
        row = await cursor.fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Client negasit.")
    return dict(row)


@router.put("/clients/{client_id}")
async def update_client(client_id: int, data: ClientUpdate):
    """Actualizare client existent."""
    async with get_db() as db:
        cursor = await db.execute(
            "SELECT * FROM clients WHERE id = ?", (client_id,)
        )
        existing = await cursor.fetchone()
        if not existing:
            raise HTTPException(status_code=404, detail="Client negasit.")

        updates = data.model_dump(exclude_unset=True)
        if not updates:
            raise HTTPException(status_code=400, detail="Niciun camp de actualizat.")

        set_clause = ", ".join(f"{k} = ?" for k in updates)
        values = list(updates.values())
        values.append(client_id)

        await db.execute(
            f"UPDATE clients SET {set_clause}, updated_at = CURRENT_TIMESTAMP WHERE id = ?",
            values,
        )
        await db.commit()

    await log_activity(
        action="invoice.client_update",
        summary=f"Client actualizat: ID {client_id}",
        details={"client_id": client_id, "fields": list(updates.keys())},
    )
    return {"message": "Client actualizat cu succes."}


@router.delete("/clients/{client_id}")
async def delete_client(client_id: int):
    """Stergere client (doar daca nu are facturi)."""
    async with get_db() as db:
        cursor = await db.execute(
            "SELECT COUNT(*) as cnt FROM invoices WHERE client_id = ?", (client_id,)
        )
        row = await cursor.fetchone()
        if row and row["cnt"] > 0:
            raise HTTPException(
                status_code=409,
                detail=f"Clientul are {row['cnt']} facturi asociate. Sterge facturile intai.",
            )

        cursor = await db.execute(
            "SELECT name FROM clients WHERE id = ?", (client_id,)
        )
        client_row = await cursor.fetchone()
        if not client_row:
            raise HTTPException(status_code=404, detail="Client negasit.")
        client_name = client_row["name"]

        await db.execute("DELETE FROM clients WHERE id = ?", (client_id,))
        await db.commit()

    await log_activity(
        action="invoice.client_delete",
        summary=f"Client sters: {client_name}",
        details={"client_id": client_id},
    )
    return {"message": f"Client '{client_name}' sters cu succes."}


# ═══════════════════════════════════════════
# Invoice endpoints
# ═══════════════════════════════════════════

@router.get("/list")
async def list_invoices():
    """Lista toate facturile cu numele clientului."""
    async with get_db() as db:
        cursor = await db.execute(
            """SELECT i.*, c.name as client_name
               FROM invoices i
               LEFT JOIN clients c ON i.client_id = c.id
               ORDER BY i.date DESC, i.id DESC"""
        )
        rows = await cursor.fetchall()
    return [_row_to_dict(row) for row in rows]


@router.post("/create", status_code=201)
async def create_invoice(data: InvoiceCreate):
    """Creare factura noua cu calcul automat subtotal/TVA/total."""
    # Validate client exists
    async with get_db() as db:
        cursor = await db.execute(
            "SELECT id, name FROM clients WHERE id = ?", (data.client_id,)
        )
        client = await cursor.fetchone()
        if not client:
            raise HTTPException(status_code=404, detail="Client negasit.")

    # Recalculate item totals and invoice totals
    for item in data.items:
        item.total = round(item.quantity * item.unit_price, 2)

    subtotal, vat_amount, total = _calculate_totals(data.items, data.vat_percent)
    invoice_number = await _next_invoice_number(data.series_prefix)
    items_json = _items_to_json(data.items)

    async with get_db() as db:
        cursor = await db.execute(
            """INSERT INTO invoices
               (client_id, invoice_number, date, due_date, items_json,
                subtotal, vat_percent, vat_amount, total, notes, status)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'draft')""",
            (
                data.client_id, invoice_number, data.date, data.due_date,
                items_json, subtotal, data.vat_percent, vat_amount, total, data.notes,
            ),
        )
        await db.commit()
        invoice_id = cursor.lastrowid

    await log_activity(
        action="invoice.create",
        summary=f"Factura {invoice_number} creata — {total} RON",
        details={
            "invoice_id": invoice_id,
            "invoice_number": invoice_number,
            "client": client["name"],
            "total": total,
        },
    )
    return {
        "id": invoice_id,
        "invoice_number": invoice_number,
        "total": total,
        "message": f"Factura {invoice_number} creata cu succes.",
    }



# ═══════════════════════════════════════════
# 10.4 — Raport lunar venituri
# (MUST be defined before /{invoice_id} to avoid path conflict)
# ═══════════════════════════════════════════

@router.get("/reports/monthly")
async def monthly_report(year: int = None):
    """Raport lunar venituri — suma per luna, pentru grafice."""
    if year is None:
        year = date.today().year

    async with get_db() as db:
        cursor = await db.execute(
            """SELECT
                 strftime('%m', date) as luna,
                 COUNT(*) as nr_facturi,
                 COALESCE(SUM(total), 0) as total_venituri,
                 COALESCE(SUM(CASE WHEN status = 'paid' THEN total ELSE 0 END), 0) as total_incasat
               FROM invoices
               WHERE strftime('%Y', date) = ?
               GROUP BY strftime('%m', date)
               ORDER BY luna ASC""",
            (str(year),),
        )
        rows = await cursor.fetchall()

    # Genereaza toate cele 12 luni (si cele fara facturi)
    month_names = [
        "Ianuarie", "Februarie", "Martie", "Aprilie",
        "Mai", "Iunie", "Iulie", "August",
        "Septembrie", "Octombrie", "Noiembrie", "Decembrie",
    ]
    data_by_month = {row["luna"]: dict(row) for row in rows}

    months = []
    for i in range(1, 13):
        key = f"{i:02d}"
        entry = data_by_month.get(key, {})
        months.append({
            "luna": i,
            "luna_nume": month_names[i - 1],
            "nr_facturi": entry.get("nr_facturi", 0),
            "total_venituri": round(entry.get("total_venituri", 0), 2),
            "total_incasat": round(entry.get("total_incasat", 0), 2),
        })

    return {
        "year": year,
        "months": months,
        "total_anual": round(sum(m["total_venituri"] for m in months), 2),
        "total_incasat_anual": round(sum(m["total_incasat"] for m in months), 2),
    }


@router.get("/reports/summary")
async def reports_summary():
    """Sumar total: venituri, medie factura, top clienti, status breakdown."""
    async with get_db() as db:
        # Total general
        cursor = await db.execute(
            """SELECT
                 COUNT(*) as total_facturi,
                 COALESCE(SUM(total), 0) as total_venituri,
                 COALESCE(AVG(total), 0) as medie_factura,
                 COALESCE(MIN(total), 0) as min_factura,
                 COALESCE(MAX(total), 0) as max_factura,
                 COALESCE(SUM(CASE WHEN status = 'paid' THEN total ELSE 0 END), 0) as total_incasat
               FROM invoices"""
        )
        totals = await cursor.fetchone()

        # Status breakdown
        cursor = await db.execute(
            """SELECT status, COUNT(*) as count, COALESCE(SUM(total), 0) as total
               FROM invoices
               GROUP BY status
               ORDER BY count DESC"""
        )
        status_rows = await cursor.fetchall()

        # Top 10 clienti dupa valoare
        cursor = await db.execute(
            """SELECT c.id, c.name, c.cui,
                      COUNT(i.id) as nr_facturi,
                      COALESCE(SUM(i.total), 0) as total_valoare
               FROM clients c
               LEFT JOIN invoices i ON c.id = i.client_id
               GROUP BY c.id
               HAVING nr_facturi > 0
               ORDER BY total_valoare DESC
               LIMIT 10"""
        )
        top_clients = await cursor.fetchall()

        # Facturi recente
        cursor = await db.execute(
            """SELECT i.id, i.invoice_number, i.date, i.total, i.status,
                      c.name as client_name
               FROM invoices i
               LEFT JOIN clients c ON i.client_id = c.id
               ORDER BY i.date DESC, i.id DESC
               LIMIT 5"""
        )
        recent = await cursor.fetchall()

    return {
        "totals": {
            "total_facturi": totals["total_facturi"],
            "total_venituri": round(totals["total_venituri"], 2),
            "medie_factura": round(totals["medie_factura"], 2),
            "min_factura": round(totals["min_factura"], 2),
            "max_factura": round(totals["max_factura"], 2),
            "total_incasat": round(totals["total_incasat"], 2),
        },
        "status_breakdown": [
            {"status": row["status"], "count": row["count"], "total": round(row["total"], 2)}
            for row in status_rows
        ],
        "top_clients": [
            {
                "id": row["id"],
                "name": row["name"],
                "cui": row["cui"],
                "nr_facturi": row["nr_facturi"],
                "total_valoare": round(row["total_valoare"], 2),
            }
            for row in top_clients
        ],
        "recent_invoices": [dict(row) for row in recent],
    }


# ═══════════════════════════════════════════
# 10.5 — Export rapoarte Excel/CSV
# (MUST be defined before /{invoice_id} to avoid path conflict)
# ═══════════════════════════════════════════

@router.get("/export/csv")
async def export_csv():
    """Export toate facturile ca fisier CSV."""
    async with get_db() as db:
        cursor = await db.execute(
            """SELECT i.id, i.invoice_number, i.date, i.due_date, i.subtotal,
                      i.vat_percent, i.vat_amount, i.total, i.currency, i.status,
                      i.notes, i.created_at,
                      c.name as client_name, c.cui as client_cui
               FROM invoices i
               LEFT JOIN clients c ON i.client_id = c.id
               ORDER BY i.date DESC, i.id DESC"""
        )
        rows = await cursor.fetchall()

    output = io.StringIO()
    writer = csv.writer(output, delimiter=";", quoting=csv.QUOTE_MINIMAL)

    # Header
    writer.writerow([
        "ID", "Nr. Factura", "Data", "Scadenta", "Client", "CUI Client",
        "Subtotal", "TVA %", "TVA Suma", "Total", "Moneda", "Status", "Observatii", "Creat la",
    ])

    # Rows
    for row in rows:
        writer.writerow([
            row["id"], row["invoice_number"], row["date"], row["due_date"] or "",
            row["client_name"] or "", row["client_cui"] or "",
            row["subtotal"], row["vat_percent"], row["vat_amount"],
            row["total"], row["currency"], row["status"],
            row["notes"] or "", row["created_at"],
        ])

    content = output.getvalue()
    output.close()

    # BOM for Excel UTF-8 compatibility
    bom = "\ufeff"
    buffer = io.BytesIO((bom + content).encode("utf-8"))

    await log_activity(
        action="invoice.export_csv",
        summary=f"Export CSV: {len(rows)} facturi",
        details={"count": len(rows)},
    )

    return StreamingResponse(
        buffer,
        media_type="text/csv; charset=utf-8",
        headers={
            "Content-Disposition": f'attachment; filename="facturi_export_{date.today().isoformat()}.csv"',
        },
    )


@router.get("/export/excel")
async def export_excel():
    """Export toate facturile ca fisier Excel (.xlsx)."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    async with get_db() as db:
        cursor = await db.execute(
            """SELECT i.id, i.invoice_number, i.date, i.due_date, i.subtotal,
                      i.vat_percent, i.vat_amount, i.total, i.currency, i.status,
                      i.notes, i.created_at,
                      c.name as client_name, c.cui as client_cui
               FROM invoices i
               LEFT JOIN clients c ON i.client_id = c.id
               ORDER BY i.date DESC, i.id DESC"""
        )
        rows = await cursor.fetchall()

    wb = Workbook()
    ws = wb.active
    ws.title = "Facturi"

    # Styles
    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_fill = PatternFill(start_color="1a365d", end_color="1a365d", fill_type="solid")
    header_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    thin_border = Border(
        left=Side(style="thin", color="CBD5E0"),
        right=Side(style="thin", color="CBD5E0"),
        top=Side(style="thin", color="CBD5E0"),
        bottom=Side(style="thin", color="CBD5E0"),
    )

    headers = [
        "ID", "Nr. Factura", "Data", "Scadenta", "Client", "CUI Client",
        "Subtotal", "TVA %", "TVA Suma", "Total", "Moneda", "Status", "Observatii", "Creat la",
    ]

    # Write header
    for col_idx, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_idx, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align
        cell.border = thin_border

    # Write data
    for row_idx, row in enumerate(rows, 2):
        values = [
            row["id"], row["invoice_number"], row["date"], row["due_date"] or "",
            row["client_name"] or "", row["client_cui"] or "",
            row["subtotal"], row["vat_percent"], row["vat_amount"],
            row["total"], row["currency"], row["status"],
            row["notes"] or "", row["created_at"],
        ]
        for col_idx, val in enumerate(values, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=val)
            cell.border = thin_border

    # Column widths
    col_widths = [6, 18, 12, 12, 25, 15, 12, 8, 12, 12, 8, 10, 30, 20]
    for i, width in enumerate(col_widths, 1):
        ws.column_dimensions[ws.cell(row=1, column=i).column_letter].width = width

    # Auto-filter
    ws.auto_filter.ref = f"A1:N{len(rows) + 1}"

    # Total row
    total_row = len(rows) + 2
    ws.cell(row=total_row, column=5, value="TOTAL:").font = Font(bold=True, size=11)
    ws.cell(row=total_row, column=10, value=sum(r["total"] for r in rows)).font = Font(bold=True, size=11)

    # Save to buffer
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)

    await log_activity(
        action="invoice.export_excel",
        summary=f"Export Excel: {len(rows)} facturi",
        details={"count": len(rows)},
    )

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={
            "Content-Disposition": f'attachment; filename="facturi_export_{date.today().isoformat()}.xlsx"',
        },
    )


# ═══════════════════════════════════════════
# 10.6 — Template documente
# (MUST be defined before /{invoice_id} to avoid path conflict)
# ═══════════════════════════════════════════

TEMPLATES_DIR = settings.data_dir / "templates"
TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)


@router.get("/templates")
async def list_templates():
    """Lista template-uri documente disponibile."""
    async with get_db() as db:
        cursor = await db.execute(
            "SELECT * FROM document_templates ORDER BY template_type, name"
        )
        rows = await cursor.fetchall()

    result = []
    for row in rows:
        entry = dict(row)
        try:
            entry["fields"] = json.loads(entry.pop("fields_json"))
        except (json.JSONDecodeError, KeyError):
            entry["fields"] = []
        result.append(entry)

    return result


@router.post("/templates/generate")
async def generate_template(data: TemplateGenerateRequest):
    """Genereaza document din template (contract, oferta, chitanta)."""
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Fetch template
    async with get_db() as db:
        cursor = await db.execute(
            "SELECT * FROM document_templates WHERE name = ?",
            (data.template_name,),
        )
        template = await cursor.fetchone()

    if not template:
        raise HTTPException(status_code=404, detail="Template negasit.")

    template = dict(template)
    template_type = template["template_type"]

    # Build document based on template type
    doc = DocxDocument()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    today_str = data.data.get("date", date.today().isoformat())

    if template_type == "contract":
        _build_contract_doc(doc, data.data, today_str)
    elif template_type == "offer":
        _build_offer_doc(doc, data.data, today_str)
    elif template_type == "receipt":
        _build_receipt_doc(doc, data.data, today_str)
    else:
        raise HTTPException(status_code=400, detail=f"Tip template necunoscut: {template_type}")

    # Save DOCX
    filename = f"{data.template_name}_{today_str}"
    docx_path = TEMPLATES_DIR / f"{filename}.docx"
    doc.save(str(docx_path))

    await log_activity(
        action="invoice.template_generate",
        summary=f"Document generat: {template['display_name']}",
        details={
            "template": data.template_name,
            "type": template_type,
            "filename": f"{filename}.docx",
        },
    )

    return FileResponse(
        path=str(docx_path),
        filename=f"{filename}.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


# ═══════════════════════════════════════════
# 10.8 — Scanner facturi primite (OCR + AI)
# (MUST be defined before /{invoice_id} to avoid path conflict)
# ═══════════════════════════════════════════

SCANNED_DIR = settings.data_dir / "scanned_invoices"
SCANNED_DIR.mkdir(parents=True, exist_ok=True)


@router.post("/scan")
async def scan_received_invoice(file: UploadFile = File(...)):
    """
    Upload factura primita (imagine/PDF), OCR + AI extrage:
    furnizor, suma, data, numar factura. Salveaza in received_invoices.
    """
    # Save uploaded file
    filename = file.filename or "scan_unknown"
    safe_name = f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{filename}"
    file_path = SCANNED_DIR / safe_name

    content = await file.read()
    with open(file_path, "wb") as f:
        f.write(content)

    # OCR text extraction
    raw_text = ""
    fname_lower = filename.lower()

    try:
        if fname_lower.endswith(".pdf"):
            import fitz
            doc = fitz.open(str(file_path))
            for page in doc:
                page_text = page.get_text()
                raw_text += page_text
                # If no text, try OCR on page image
                if not page_text.strip():
                    pix = page.get_pixmap(dpi=300)
                    img_path = str(file_path) + f"_page.png"
                    pix.save(img_path)
                    try:
                        import pytesseract
                        from PIL import Image
                        img = Image.open(img_path)
                        raw_text += pytesseract.image_to_string(img, lang="ron+eng")
                    except ImportError:
                        raw_text += "[pytesseract indisponibil]"
                    finally:
                        if os.path.exists(img_path):
                            os.unlink(img_path)
            doc.close()
        elif fname_lower.endswith((".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".webp")):
            try:
                import pytesseract
                from PIL import Image
                img = Image.open(str(file_path))
                raw_text = pytesseract.image_to_string(img, lang="ron+eng")
            except ImportError:
                raise HTTPException(500, "pytesseract nu e instalat")
        else:
            raise HTTPException(
                400,
                "Format nesuportat. Acceptate: PDF, PNG, JPG, TIFF, BMP, WEBP",
            )
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("Eroare OCR: %s", exc)
        raise HTTPException(500, f"Eroare la procesarea fisierului: {exc}")

    if not raw_text.strip():
        # Save even without text
        async with get_db() as db:
            cursor = await db.execute(
                """INSERT INTO received_invoices (file_path, raw_text, extracted_data)
                   VALUES (?, ?, ?)""",
                (str(file_path), "", "{}"),
            )
            await db.commit()
            scan_id = cursor.lastrowid

        return {
            "id": scan_id,
            "message": "Nu s-a detectat text in document.",
            "raw_text": "",
            "extracted": {},
        }

    # AI extraction of invoice data
    extracted = {}
    try:
        from modules.ai.providers import ai_generate

        prompt = f"""Analizeaza textul urmator extras dintr-o factura primita si extrage urmatoarele informatii.
Returneaza un JSON valid cu campurile gasite (omite ce nu exista):
- supplier_name: numele furnizorului/emitentului
- supplier_cui: CUI furnizor
- invoice_number: numarul facturii
- invoice_date: data facturii (format YYYY-MM-DD daca posibil)
- amount: suma fara TVA (doar numarul)
- vat: suma TVA (doar numarul)
- total: total de plata (doar numarul)
- currency: moneda (RON, EUR, USD)

Returneaza DOAR JSON valid, fara explicatii.
Daca nu gasesti informatii, returneaza: {{}}

Text factura:
{raw_text[:3000]}"""

        system_prompt = "Esti un asistent care extrage date structurate din facturi. Raspunzi DOAR cu JSON valid."
        result = await ai_generate(prompt, system_prompt=system_prompt)
        response_text = result.get("text", "").strip()

        # Clean markdown code block markers
        if response_text.startswith("```"):
            lines = response_text.split("\n")
            lines = [l for l in lines if not l.strip().startswith("```")]
            response_text = "\n".join(lines).strip()

        parsed = json.loads(response_text)
        if isinstance(parsed, dict):
            allowed = {"supplier_name", "supplier_cui", "invoice_number", "invoice_date",
                       "amount", "vat", "total", "currency"}
            extracted = {k: v for k, v in parsed.items() if k in allowed and v is not None}
    except (json.JSONDecodeError, RuntimeError, Exception) as exc:
        logger.warning("AI extraction pentru factura scanata a esuat: %s", exc)

    # Save to DB
    async with get_db() as db:
        cursor = await db.execute(
            """INSERT INTO received_invoices
               (supplier_name, supplier_cui, invoice_number, invoice_date,
                amount, vat, total, currency, file_path, raw_text, extracted_data)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (
                extracted.get("supplier_name"),
                extracted.get("supplier_cui"),
                extracted.get("invoice_number"),
                extracted.get("invoice_date"),
                float(extracted["amount"]) if extracted.get("amount") else 0,
                float(extracted["vat"]) if extracted.get("vat") else 0,
                float(extracted["total"]) if extracted.get("total") else 0,
                extracted.get("currency", "RON"),
                str(file_path),
                raw_text,
                json.dumps(extracted, ensure_ascii=False),
            ),
        )
        await db.commit()
        scan_id = cursor.lastrowid

    await log_activity(
        action="invoice.scan",
        summary=f"Factura scanata: {extracted.get('supplier_name', 'necunoscut')} — {extracted.get('total', '?')} {extracted.get('currency', 'RON')}",
        details={
            "scan_id": scan_id,
            "filename": filename,
            "extracted": extracted,
        },
    )

    return {
        "id": scan_id,
        "message": "Factura scanata si procesata cu succes.",
        "raw_text": raw_text.strip()[:2000],
        "extracted": extracted,
    }


# ═══════════════════════════════════════════
# F3: Invoice Series CRUD (MUST be before /{invoice_id})
# ═══════════════════════════════════════════

@router.get("/series")
async def list_series_early():
    """Lista toate seriile de facturi."""
    async with get_db() as db:
        cursor = await db.execute("SELECT * FROM invoice_series ORDER BY is_default DESC, name ASC")
        rows = await cursor.fetchall()
    return [dict(r) for r in rows]


@router.post("/series", status_code=201)
async def create_series_early(data: InvoiceSeriesCreate):
    """Creare serie noua de facturi."""
    async with get_db() as db:
        cursor = await db.execute(
            "INSERT INTO invoice_series (prefix, name, description) VALUES (?, ?, ?)",
            (data.prefix.upper().strip(), data.name, data.description),
        )
        await db.commit()
        series_id = cursor.lastrowid
    await log_activity(
        action="invoice.series_create",
        summary=f"Serie facturi creata: {data.prefix} — {data.name}",
    )
    return {"id": series_id, "message": f"Serie '{data.prefix}' creata cu succes."}


@router.put("/series/{series_id}")
async def update_series_early(series_id: int, data: InvoiceSeriesUpdate):
    """Actualizare serie facturi."""
    updates = data.model_dump(exclude_unset=True)
    if not updates:
        raise HTTPException(400, "Niciun camp de actualizat")
    if "active" in updates:
        updates["active"] = 1 if updates["active"] else 0
    fields = ", ".join(f"{k} = ?" for k in updates)
    values = list(updates.values()) + [series_id]
    async with get_db() as db:
        cursor = await db.execute(
            f"UPDATE invoice_series SET {fields} WHERE id = ?", values
        )
        await db.commit()
        if cursor.rowcount == 0:
            raise HTTPException(404, "Serie negasita")
    return {"message": "Serie actualizata cu succes."}


@router.put("/series/{series_id}/default")
async def set_default_series_early(series_id: int):
    """Seteaza o serie ca implicita."""
    async with get_db() as db:
        await db.execute("UPDATE invoice_series SET is_default = 0")
        cursor = await db.execute(
            "UPDATE invoice_series SET is_default = 1 WHERE id = ?", (series_id,)
        )
        await db.commit()
        if cursor.rowcount == 0:
            raise HTTPException(404, "Serie negasita")
    return {"message": "Serie setata ca implicita."}


@router.delete("/series/{series_id}")
async def delete_series_early(series_id: int):
    """Sterge o serie (doar daca nu e implicita si nu are facturi)."""
    async with get_db() as db:
        cursor = await db.execute(
            "SELECT is_default, prefix FROM invoice_series WHERE id = ?", (series_id,)
        )
        row = await cursor.fetchone()
        if not row:
            raise HTTPException(404, "Serie negasita")
        if row["is_default"]:
            raise HTTPException(409, "Nu se poate sterge seria implicita")
        cursor = await db.execute(
            "SELECT COUNT(*) as cnt FROM invoices WHERE series = ?", (row["prefix"],)
        )
        cnt_row = await cursor.fetchone()
        if cnt_row and cnt_row["cnt"] > 0:
            raise HTTPException(409, f"Seria are {cnt_row['cnt']} facturi asociate")
        await db.execute("DELETE FROM invoice_series WHERE id = ?", (series_id,))
        await db.commit()
    return {"message": "Serie stearsa cu succes."}


# ═══════════════════════════════════════════
# F4: Overdue invoices (MUST be before /{invoice_id})
# ═══════════════════════════════════════════

@router.get("/overdue")
async def list_overdue_early():
    """Lista facturi cu scadenta depasita (neplatite)."""
    today = date.today().isoformat()
    async with get_db() as db:
        cursor = await db.execute(
            """SELECT i.*, c.name as client_name, c.email as client_email, c.phone as client_phone
               FROM invoices i
               LEFT JOIN clients c ON i.client_id = c.id
               WHERE i.due_date IS NOT NULL AND i.due_date < ? AND i.status NOT IN ('paid', 'cancelled')
               ORDER BY i.due_date ASC""",
            (today,),
        )
        rows = await cursor.fetchall()
    result = []
    today_d = date.today()
    for row in rows:
        d = _row_to_dict(row)
        try:
            due = date.fromisoformat(d["due_date"])
            d["days_overdue"] = (today_d - due).days
        except (ValueError, TypeError):
            d["days_overdue"] = 0
        result.append(d)
    return result


# ═══════════════════════════════════════════
# F9: Offer/Quote PDF (MUST be before /{invoice_id})
# ═══════════════════════════════════════════

@router.post("/offer-pdf")
async def generate_offer_pdf_early(data: OfferPdfRequest):
    """Generate an offer/quote PDF for CIP Inspection SRL."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm
        from reportlab.pdfgen import canvas as pdf_canvas
    except ImportError:
        return await _generate_offer_fitz(data)

    buf = io.BytesIO()
    c = pdf_canvas.Canvas(buf, pagesize=A4)
    width, height = A4

    c.setFont("Helvetica-Bold", 16)
    c.drawString(2 * cm, height - 2 * cm, "OFERTA DE PRET")
    c.setFont("Helvetica", 10)
    c.drawString(2 * cm, height - 2.8 * cm, "CIP Inspection SRL | CUI 43978110")
    c.drawString(2 * cm, height - 3.4 * cm, f"Data: {date.today().strftime('%d.%m.%Y')}")
    c.drawString(2 * cm, height - 3.9 * cm, f"Valabilitate: {data.validity_days} zile")

    c.setFont("Helvetica-Bold", 11)
    c.drawString(2 * cm, height - 5 * cm, f"Catre: {data.client_name}")
    if data.client_address:
        c.setFont("Helvetica", 10)
        c.drawString(2 * cm, height - 5.5 * cm, data.client_address)

    y = height - 7 * cm
    c.setFont("Helvetica-Bold", 10)
    c.drawString(2 * cm, y, "Nr.")
    c.drawString(3 * cm, y, "Descriere")
    c.drawString(12 * cm, y, "Cant.")
    c.drawString(14 * cm, y, "Pret unit.")
    c.drawString(17 * cm, y, "Total")
    y -= 0.6 * cm

    c.setFont("Helvetica", 10)
    subtotal = 0
    for i, item in enumerate(data.items, 1):
        total_item = round(item.quantity * item.unit_price, 2)
        subtotal += total_item
        c.drawString(2 * cm, y, str(i))
        c.drawString(3 * cm, y, item.description[:50])
        c.drawString(12 * cm, y, str(item.quantity))
        c.drawString(14 * cm, y, f"{item.unit_price:.2f}")
        c.drawString(17 * cm, y, f"{total_item:.2f}")
        y -= 0.5 * cm

    y -= 0.5 * cm
    c.setFont("Helvetica-Bold", 11)
    c.drawString(14 * cm, y, f"TOTAL: {subtotal:.2f} RON")

    if data.notes:
        y -= 1.5 * cm
        c.setFont("Helvetica", 9)
        c.drawString(2 * cm, y, f"Observatii: {data.notes}")

    c.save()
    buf.seek(0)

    await log_activity(
        action="invoice.offer_pdf",
        summary=f"Oferta PDF generata: {data.client_name} — {subtotal:.2f} RON",
        details={"client": data.client_name, "total": subtotal, "items": len(data.items)},
    )

    filename = f"oferta_{data.client_name.replace(' ', '_')}_{date.today().isoformat()}.pdf"
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


async def _generate_offer_fitz(data: OfferPdfRequest):
    """Fallback offer PDF using PyMuPDF (fitz) if reportlab not available."""
    import fitz as fitz_mod

    doc = fitz_mod.open()
    page = doc.new_page()

    y = 50
    page.insert_text((50, y), "OFERTA DE PRET", fontsize=18, fontname="helv")
    y += 25
    page.insert_text((50, y), "CIP Inspection SRL | CUI 43978110", fontsize=10)
    y += 15
    page.insert_text((50, y), f"Data: {date.today().strftime('%d.%m.%Y')} | Valabilitate: {data.validity_days} zile", fontsize=10)
    y += 25
    page.insert_text((50, y), f"Catre: {data.client_name}", fontsize=12, fontname="helv")
    if data.client_address:
        y += 15
        page.insert_text((50, y), data.client_address, fontsize=10)
    y += 30

    subtotal = 0
    for i, item in enumerate(data.items, 1):
        total_item = round(item.quantity * item.unit_price, 2)
        subtotal += total_item
        line = f"{i}. {item.description} — {item.quantity} x {item.unit_price:.2f} = {total_item:.2f} RON"
        page.insert_text((50, y), line, fontsize=10)
        y += 15

    y += 10
    page.insert_text((50, y), f"TOTAL: {subtotal:.2f} RON", fontsize=13, fontname="helv")

    if data.notes:
        y += 25
        page.insert_text((50, y), f"Observatii: {data.notes}", fontsize=9)

    buf = io.BytesIO()
    doc.save(buf)
    doc.close()
    buf.seek(0)

    await log_activity(
        action="invoice.offer_pdf",
        summary=f"Oferta PDF (fitz): {data.client_name} — {subtotal:.2f} RON",
    )

    filename = f"oferta_{data.client_name.replace(' ', '_')}_{date.today().isoformat()}.pdf"
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ═══════════════════════════════════════════
# Invoice detail endpoints (parametric /{invoice_id})
# ═══════════════════════════════════════════

@router.get("/{invoice_id}")
async def get_invoice(invoice_id: int):
    """Detalii factura cu informatii client."""
    async with get_db() as db:
        cursor = await db.execute(
            """SELECT i.*, c.name as client_name, c.cui as client_cui,
                      c.reg_com as client_reg_com, c.address as client_address,
                      c.email as client_email, c.phone as client_phone
               FROM invoices i
               LEFT JOIN clients c ON i.client_id = c.id
               WHERE i.id = ?""",
            (invoice_id,),
        )
        row = await cursor.fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Factura negasita.")
    return _row_to_dict(row)


@router.put("/{invoice_id}")
async def update_invoice(invoice_id: int, data: InvoiceUpdate):
    """Actualizare factura (doar daca este in status draft)."""
    async with get_db() as db:
        cursor = await db.execute(
            "SELECT * FROM invoices WHERE id = ?", (invoice_id,)
        )
        existing = await cursor.fetchone()
        if not existing:
            raise HTTPException(status_code=404, detail="Factura negasita.")
        if existing["status"] != "draft":
            raise HTTPException(
                status_code=409,
                detail=f"Factura nu poate fi modificata — status curent: {existing['status']}.",
            )

        updates = data.model_dump(exclude_unset=True)
        if not updates:
            raise HTTPException(status_code=400, detail="Niciun camp de actualizat.")

        # Validate new client if provided
        if "client_id" in updates:
            cursor = await db.execute(
                "SELECT id FROM clients WHERE id = ?", (updates["client_id"],)
            )
            if not await cursor.fetchone():
                raise HTTPException(status_code=404, detail="Client negasit.")

        # Recalculate if items or vat changed
        if "items" in updates:
            items = [InvoiceItem(**it) for it in updates["items"]] if isinstance(updates["items"][0], dict) else updates["items"]
            for item in items:
                item.total = round(item.quantity * item.unit_price, 2)
            vat_pct = updates.get("vat_percent", existing["vat_percent"])
            subtotal, vat_amount, total = _calculate_totals(items, vat_pct)
            updates["items_json"] = _items_to_json(items)
            updates["subtotal"] = subtotal
            updates["vat_amount"] = vat_amount
            updates["total"] = total
            del updates["items"]
        elif "vat_percent" in updates:
            # Recalculate with existing items
            old_items = _items_from_json(existing["items_json"])
            items = [InvoiceItem(**it) for it in old_items]
            subtotal, vat_amount, total = _calculate_totals(items, updates["vat_percent"])
            updates["subtotal"] = subtotal
            updates["vat_amount"] = vat_amount
            updates["total"] = total

        set_clause = ", ".join(f"{k} = ?" for k in updates)
        values = list(updates.values())
        values.append(invoice_id)

        await db.execute(
            f"UPDATE invoices SET {set_clause}, updated_at = CURRENT_TIMESTAMP WHERE id = ?",
            values,
        )
        await db.commit()

    await log_activity(
        action="invoice.update",
        summary=f"Factura ID {invoice_id} actualizata",
        details={"invoice_id": invoice_id, "fields": list(data.model_dump(exclude_unset=True).keys())},
    )
    return {"message": "Factura actualizata cu succes."}


@router.delete("/{invoice_id}")
async def delete_invoice(invoice_id: int):
    """Stergere factura (doar daca este in status draft)."""
    async with get_db() as db:
        cursor = await db.execute(
            "SELECT invoice_number, status, pdf_path FROM invoices WHERE id = ?",
            (invoice_id,),
        )
        row = await cursor.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Factura negasita.")
        if row["status"] != "draft":
            raise HTTPException(
                status_code=409,
                detail=f"Factura nu poate fi stearsa — status curent: {row['status']}.",
            )
        invoice_number = row["invoice_number"]

        # Delete PDF file if it exists
        if row["pdf_path"]:
            pdf_path = Path(row["pdf_path"])
            if pdf_path.exists():
                pdf_path.unlink()

        await db.execute("DELETE FROM invoices WHERE id = ?", (invoice_id,))
        await db.commit()

    await log_activity(
        action="invoice.delete",
        summary=f"Factura {invoice_number} stearsa",
        details={"invoice_id": invoice_id, "invoice_number": invoice_number},
    )
    return {"message": f"Factura {invoice_number} stearsa cu succes."}


@router.post("/{invoice_id}/pdf")
async def generate_pdf(invoice_id: int):
    """Genereaza PDF profesional si returneaza fisierul."""
    # Fetch invoice + client
    async with get_db() as db:
        cursor = await db.execute(
            """SELECT i.*, c.name as client_name, c.cui as client_cui,
                      c.reg_com as client_reg_com, c.address as client_address,
                      c.email as client_email, c.phone as client_phone
               FROM invoices i
               LEFT JOIN clients c ON i.client_id = c.id
               WHERE i.id = ?""",
            (invoice_id,),
        )
        row = await cursor.fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Factura negasita.")

    invoice = dict(row)
    items = _items_from_json(invoice.get("items_json", "[]"))

    pdf_filename = f"{invoice['invoice_number']}.pdf"
    pdf_path = INVOICES_DIR / pdf_filename

    try:
        _build_invoice_pdf(invoice, items, str(pdf_path))
    except Exception as exc:
        logger.error("Eroare la generare PDF: %s", exc)
        raise HTTPException(status_code=500, detail=f"Eroare la generare PDF: {exc}")

    # Save path in DB
    async with get_db() as db:
        await db.execute(
            "UPDATE invoices SET pdf_path = ?, updated_at = CURRENT_TIMESTAMP WHERE id = ?",
            (str(pdf_path), invoice_id),
        )
        await db.commit()

    await log_activity(
        action="invoice.pdf",
        summary=f"PDF generat: {invoice['invoice_number']}",
        details={"invoice_id": invoice_id, "pdf_path": str(pdf_path)},
    )
    return FileResponse(
        path=str(pdf_path),
        filename=pdf_filename,
        media_type="application/pdf",
    )


@router.put("/{invoice_id}/status")
async def change_status(invoice_id: int, data: StatusUpdate):
    """Schimbare status factura cu validari de tranzitie."""
    valid_transitions = {
        "draft": ["sent", "cancelled"],
        "sent": ["paid", "cancelled"],
        "paid": [],
        "cancelled": ["draft"],
    }

    async with get_db() as db:
        cursor = await db.execute(
            "SELECT invoice_number, status FROM invoices WHERE id = ?",
            (invoice_id,),
        )
        row = await cursor.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Factura negasita.")

        current = row["status"]
        new_status = data.status.lower()

        if new_status not in valid_transitions.get(current, []):
            allowed = valid_transitions.get(current, [])
            raise HTTPException(
                status_code=409,
                detail=f"Tranzitie invalida: {current} -> {new_status}. Tranzitii permise: {allowed or 'niciuna'}.",
            )

        await db.execute(
            "UPDATE invoices SET status = ?, updated_at = CURRENT_TIMESTAMP WHERE id = ?",
            (new_status, invoice_id),
        )
        await db.commit()

    await log_activity(
        action="invoice.status_change",
        summary=f"Factura {row['invoice_number']}: {current} -> {new_status}",
        details={
            "invoice_id": invoice_id,
            "invoice_number": row["invoice_number"],
            "old_status": current,
            "new_status": new_status,
        },
    )
    return {"message": f"Status actualizat: {current} -> {new_status}."}


@router.put("/{invoice_id}/payment")
async def mark_payment(invoice_id: int):
    """Marcheaza o factura ca platita cu data curenta."""
    today = date.today().isoformat()
    async with get_db() as db:
        cursor = await db.execute(
            "UPDATE invoices SET status = 'paid', payment_date = ?, updated_at = CURRENT_TIMESTAMP WHERE id = ?",
            (today, invoice_id),
        )
        await db.commit()
        if cursor.rowcount == 0:
            raise HTTPException(404, "Factura negasita")
    await log_activity(
        action="invoice.payment",
        summary=f"Factura #{invoice_id} marcata ca platita",
        details={"invoice_id": invoice_id, "payment_date": today},
    )
    return {"message": "Factura marcata ca platita.", "payment_date": today}


# ═══════════════════════════════════════════
# AI-assisted: generate from calculation
# ═══════════════════════════════════════════

@router.post("/generate-from-calc")
async def generate_from_calc(data: GenerateFromCalcRequest):
    """
    Pre-fill date factura din calculul de pret existent.

    Citeste din uploads + calculations, foloseste AI (Gemini Flash)
    pentru a extrage informatii client din document (daca e disponibil).
    Returneaza date pre-completate — NU salveaza in DB.
    """
    # Fetch calculation + upload data
    async with get_db() as db:
        cursor = await db.execute(
            """SELECT c.*, u.filename, u.original_name, u.filepath, u.file_type
               FROM calculations c
               JOIN uploads u ON c.upload_id = u.id
               WHERE c.id = ?""",
            (data.calculation_id,),
        )
        calc_row = await cursor.fetchone()

    if not calc_row:
        raise HTTPException(
            status_code=404,
            detail="Calculul nu a fost gasit sau nu are upload asociat.",
        )

    calc = dict(calc_row)
    final_price = calc.get("final_price", 0) or calc.get("estimated_price", 0) or 0
    doc_name = calc.get("original_name") or calc.get("filename") or "Document"
    filepath = calc.get("filepath", "")

    # Build basic invoice item from calculation
    items = [{
        "description": f"Traducere document: {doc_name}",
        "quantity": 1,
        "unit_price": round(float(final_price), 2),
        "total": round(float(final_price), 2),
    }]

    # Try AI extraction of client info from document
    client_suggestion: dict[str, Any] = {}
    if filepath and Path(filepath).exists():
        try:
            client_suggestion = await _extract_client_with_ai(filepath)
        except Exception as exc:
            logger.warning("AI client extraction failed: %s", exc)

    result = {
        "calculation_id": data.calculation_id,
        "suggested_client": client_suggestion,
        "items": items,
        "subtotal": items[0]["total"],
        "vat_percent": 0,
        "vat_amount": 0,
        "total": items[0]["total"],
        "date": date.today().isoformat(),
        "notes": f"Generat din calculul #{data.calculation_id} — {doc_name}",
        "source": "calculation",
    }

    await log_activity(
        action="invoice.generate_from_calc",
        summary=f"Pre-fill factura din calcul #{data.calculation_id}",
        details={"calculation_id": data.calculation_id, "total": result["total"]},
    )
    return result


async def _extract_client_with_ai(filepath: str) -> dict[str, Any]:
    """
    Foloseste AI (Gemini Flash) pentru a extrage informatii client din document.
    Returneaza dict cu campuri: name, cui, address, email, phone (ce gaseste).
    """
    from modules.ai.providers import ai_generate

    # Extract text from document (first 2000 chars for efficiency)
    text = _extract_text_safe(filepath, max_chars=2000)
    if not text:
        return {}

    prompt = f"""Analizeaza textul urmator si extrage informatii despre client/beneficiar.
Returneaza un JSON valid cu campurile gasite (omite ce nu exista):
- name: numele companiei sau persoanei
- cui: codul unic de inregistrare (daca exista)
- address: adresa
- email: adresa de email
- phone: numar de telefon

Returneaza DOAR JSON valid, fara explicatii.
Daca nu gasesti informatii despre client, returneaza: {{}}

Text document:
{text}"""

    system_prompt = "Esti un asistent care extrage date structurate din documente. Raspunzi DOAR cu JSON valid."

    try:
        result = await ai_generate(prompt, system_prompt=system_prompt)
        response_text = result.get("text", "").strip()

        # Clean response — remove markdown code block markers if present
        if response_text.startswith("```"):
            lines = response_text.split("\n")
            lines = [l for l in lines if not l.strip().startswith("```")]
            response_text = "\n".join(lines).strip()

        parsed = json.loads(response_text)
        if isinstance(parsed, dict):
            # Only keep known fields
            allowed = {"name", "cui", "address", "email", "phone"}
            return {k: v for k, v in parsed.items() if k in allowed and v}
    except (json.JSONDecodeError, RuntimeError) as exc:
        logger.warning("AI client extraction parse error: %s", exc)

    return {}


def _extract_text_safe(filepath: str, max_chars: int = 2000) -> str:
    """Extract text from file (PDF, DOCX, text). Returns empty string on error."""
    try:
        p = filepath.lower()
        if p.endswith(".pdf"):
            import fitz
            doc = fitz.open(filepath)
            text = "\n".join(page.get_text() for page in doc)
            doc.close()
            return text[:max_chars].strip()
        elif p.endswith(".docx"):
            from docx import Document as DocxDocument
            doc = DocxDocument(filepath)
            text = "\n".join(para.text for para in doc.paragraphs)
            return text[:max_chars].strip()
        elif p.endswith((".txt", ".md", ".csv")):
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                return f.read(max_chars).strip()
    except Exception as exc:
        logger.warning("Text extraction failed for %s: %s", filepath, exc)
    return ""


# ═══════════════════════════════════════════
# PDF Generation with ReportLab
# ═══════════════════════════════════════════

def _build_invoice_pdf(invoice: dict, items: list[dict], output_path: str) -> None:
    """
    Genereaza PDF profesional pentru factura.

    Header: CIP Inspection SRL / CUI 43978110
    Continut: detalii client, tabel articole, totaluri
    Footer: detalii bancare placeholder
    """
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm, mm
    from reportlab.platypus import (
        SimpleDocTemplate,
        Table,
        TableStyle,
        Paragraph,
        Spacer,
    )
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        topMargin=1.5 * cm,
        bottomMargin=2 * cm,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
    )

    styles = getSampleStyleSheet()

    # Custom styles
    style_title = ParagraphStyle(
        "InvoiceTitle",
        parent=styles["Heading1"],
        fontSize=18,
        alignment=TA_CENTER,
        spaceAfter=6,
        textColor=colors.HexColor("#1a365d"),
    )
    style_company = ParagraphStyle(
        "Company",
        parent=styles["Normal"],
        fontSize=10,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#4a5568"),
    )
    style_section = ParagraphStyle(
        "Section",
        parent=styles["Heading2"],
        fontSize=11,
        spaceBefore=12,
        spaceAfter=4,
        textColor=colors.HexColor("#2d3748"),
    )
    style_normal = ParagraphStyle(
        "NormalCustom",
        parent=styles["Normal"],
        fontSize=9,
        leading=13,
    )
    style_right = ParagraphStyle(
        "RightAlign",
        parent=styles["Normal"],
        fontSize=9,
        alignment=TA_RIGHT,
    )
    style_footer = ParagraphStyle(
        "Footer",
        parent=styles["Normal"],
        fontSize=8,
        textColor=colors.HexColor("#718096"),
        alignment=TA_CENTER,
    )
    style_bold = ParagraphStyle(
        "BoldText",
        parent=styles["Normal"],
        fontSize=10,
        leading=14,
    )

    elements = []

    # --- HEADER ---
    elements.append(Paragraph("FACTURA", style_title))
    elements.append(Paragraph("CIP Inspection SRL", style_company))
    elements.append(Paragraph("CUI: 43978110", style_company))
    elements.append(Spacer(1, 8 * mm))

    # --- Invoice meta info (2-column layout) ---
    inv_number = invoice.get("invoice_number", "")
    inv_date = invoice.get("date", "")
    inv_due = invoice.get("due_date", "")
    inv_status = invoice.get("status", "draft").upper()

    meta_data = [
        [
            Paragraph(f"<b>Nr. factura:</b> {inv_number}", style_normal),
            Paragraph(f"<b>Data:</b> {inv_date}", style_normal),
        ],
        [
            Paragraph(f"<b>Status:</b> {inv_status}", style_normal),
            Paragraph(f"<b>Scadenta:</b> {inv_due or 'N/A'}", style_normal),
        ],
    ]
    meta_table = Table(meta_data, colWidths=[9 * cm, 8 * cm])
    meta_table.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    elements.append(meta_table)
    elements.append(Spacer(1, 6 * mm))

    # --- CLIENT INFO ---
    elements.append(Paragraph("Date client", style_section))
    client_name = invoice.get("client_name", "N/A")
    client_cui = invoice.get("client_cui", "")
    client_reg = invoice.get("client_reg_com", "")
    client_addr = invoice.get("client_address", "")
    client_email = invoice.get("client_email", "")
    client_phone = invoice.get("client_phone", "")

    client_lines = [f"<b>{client_name}</b>"]
    if client_cui:
        client_lines.append(f"CUI: {client_cui}")
    if client_reg:
        client_lines.append(f"Reg. Com.: {client_reg}")
    if client_addr:
        client_lines.append(f"Adresa: {client_addr}")
    if client_email:
        client_lines.append(f"Email: {client_email}")
    if client_phone:
        client_lines.append(f"Tel: {client_phone}")
    elements.append(Paragraph("<br/>".join(client_lines), style_normal))
    elements.append(Spacer(1, 6 * mm))

    # --- ITEMS TABLE ---
    elements.append(Paragraph("Articole", style_section))

    header_row = [
        Paragraph("<b>Nr.</b>", style_normal),
        Paragraph("<b>Descriere</b>", style_normal),
        Paragraph("<b>Cant.</b>", style_normal),
        Paragraph("<b>Pret unitar</b>", style_normal),
        Paragraph("<b>Total</b>", style_normal),
    ]

    table_data = [header_row]
    for idx, item in enumerate(items, 1):
        desc = item.get("description", "")
        qty = item.get("quantity", 1)
        unit_price = item.get("unit_price", 0)
        item_total = item.get("total", qty * unit_price)
        table_data.append([
            Paragraph(str(idx), style_normal),
            Paragraph(desc, style_normal),
            Paragraph(f"{qty:.2f}", style_right),
            Paragraph(f"{unit_price:.2f} RON", style_right),
            Paragraph(f"{item_total:.2f} RON", style_right),
        ])

    items_table = Table(
        table_data,
        colWidths=[1.2 * cm, 8.5 * cm, 2 * cm, 2.8 * cm, 2.8 * cm],
        repeatRows=1,
    )
    items_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#edf2f7")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#2d3748")),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5e0")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f7fafc")]),
    ]))
    elements.append(items_table)
    elements.append(Spacer(1, 4 * mm))

    # --- TOTALS ---
    subtotal = invoice.get("subtotal", 0)
    vat_pct = invoice.get("vat_percent", 0)
    vat_amount = invoice.get("vat_amount", 0)
    total = invoice.get("total", 0)

    totals_data = [
        ["", Paragraph(f"<b>Subtotal:</b>", style_right), Paragraph(f"{subtotal:.2f} RON", style_right)],
    ]
    if vat_pct > 0:
        totals_data.append(
            ["", Paragraph(f"<b>TVA ({vat_pct:.0f}%):</b>", style_right), Paragraph(f"{vat_amount:.2f} RON", style_right)]
        )
    totals_data.append(
        ["", Paragraph(f"<b>TOTAL:</b>", style_right), Paragraph(f"<b>{total:.2f} RON</b>", style_right)]
    )

    totals_table = Table(totals_data, colWidths=[10 * cm, 4 * cm, 3.3 * cm])
    totals_table.setStyle(TableStyle([
        ("LINEABOVE", (1, -1), (-1, -1), 1, colors.HexColor("#2d3748")),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    elements.append(totals_table)

    # --- NOTES ---
    notes = invoice.get("notes", "")
    if notes:
        elements.append(Spacer(1, 6 * mm))
        elements.append(Paragraph("Observatii", style_section))
        elements.append(Paragraph(notes, style_normal))

    # --- FOOTER ---
    elements.append(Spacer(1, 12 * mm))
    elements.append(Paragraph("_" * 70, style_footer))
    elements.append(Spacer(1, 2 * mm))
    elements.append(Paragraph(
        "CIP Inspection SRL | CUI: 43978110 | Cont bancar: [completati] | Banca: [completati]",
        style_footer,
    ))
    elements.append(Paragraph(
        "Factura generata automat — Roland Command Center",
        style_footer,
    ))

    doc.build(elements)


# ═══════════════════════════════════════════
# 10.3 — Istoric comenzi per client
# ═══════════════════════════════════════════

@router.get("/clients/{client_id}/history")
async def client_history(client_id: int):
    """Returneaza toate facturile unui client cu sumar totaluri."""
    async with get_db() as db:
        # Verificare client existent
        cursor = await db.execute(
            "SELECT * FROM clients WHERE id = ?", (client_id,)
        )
        client = await cursor.fetchone()
        if not client:
            raise HTTPException(status_code=404, detail="Client negasit.")

        # Facturi client
        cursor = await db.execute(
            """SELECT i.*, c.name as client_name
               FROM invoices i
               LEFT JOIN clients c ON i.client_id = c.id
               WHERE i.client_id = ?
               ORDER BY i.date DESC, i.id DESC""",
            (client_id,),
        )
        rows = await cursor.fetchall()
        invoices = [_row_to_dict(row) for row in rows]

        # Sumar totaluri
        cursor = await db.execute(
            """SELECT
                 COUNT(*) as total_facturi,
                 COALESCE(SUM(total), 0) as total_valoare,
                 COALESCE(AVG(total), 0) as medie_factura,
                 COUNT(CASE WHEN status = 'paid' THEN 1 END) as facturi_platite,
                 COUNT(CASE WHEN status = 'draft' THEN 1 END) as facturi_draft,
                 COUNT(CASE WHEN status = 'sent' THEN 1 END) as facturi_trimise,
                 COUNT(CASE WHEN status = 'cancelled' THEN 1 END) as facturi_anulate,
                 COALESCE(SUM(CASE WHEN status = 'paid' THEN total ELSE 0 END), 0) as total_incasat
               FROM invoices
               WHERE client_id = ?""",
            (client_id,),
        )
        summary_row = await cursor.fetchone()

    summary = {
        "total_facturi": summary_row["total_facturi"],
        "total_valoare": round(summary_row["total_valoare"], 2),
        "medie_factura": round(summary_row["medie_factura"], 2),
        "facturi_platite": summary_row["facturi_platite"],
        "facturi_draft": summary_row["facturi_draft"],
        "facturi_trimise": summary_row["facturi_trimise"],
        "facturi_anulate": summary_row["facturi_anulate"],
        "total_incasat": round(summary_row["total_incasat"], 2),
    }

    return {
        "client": dict(client),
        "invoices": invoices,
        "summary": summary,
    }


# ═══════════════════════════════════════════
# 10.7 — Email facturi
# ═══════════════════════════════════════════

async def _get_smtp_config() -> dict[str, str]:
    """Citeste configurarea SMTP din ai_config table."""
    keys = ["smtp_host", "smtp_port", "smtp_user", "smtp_pass", "smtp_from"]
    config = {}
    async with get_db() as db:
        for key in keys:
            cursor = await db.execute(
                "SELECT value FROM ai_config WHERE key = ?", (key,)
            )
            row = await cursor.fetchone()
            config[key] = row["value"] if row else ""
    return config


@router.post("/{invoice_id}/send-email")
async def send_invoice_email(invoice_id: int, data: SendEmailRequest):
    """Trimite factura PDF pe email."""
    # Fetch invoice
    async with get_db() as db:
        cursor = await db.execute(
            """SELECT i.*, c.name as client_name, c.cui as client_cui,
                      c.reg_com as client_reg_com, c.address as client_address,
                      c.email as client_email, c.phone as client_phone
               FROM invoices i
               LEFT JOIN clients c ON i.client_id = c.id
               WHERE i.id = ?""",
            (invoice_id,),
        )
        row = await cursor.fetchone()

    if not row:
        raise HTTPException(status_code=404, detail="Factura negasita.")

    invoice = dict(row)
    inv_number = invoice["invoice_number"]

    # Ensure PDF exists — generate if missing
    pdf_path = invoice.get("pdf_path")
    if not pdf_path or not Path(pdf_path).exists():
        items = _items_from_json(invoice.get("items_json", "[]"))
        pdf_filename = f"{inv_number}.pdf"
        pdf_path = str(INVOICES_DIR / pdf_filename)
        try:
            _build_invoice_pdf(invoice, items, pdf_path)
        except Exception as exc:
            raise HTTPException(status_code=500, detail=f"Eroare la generare PDF: {exc}")

        # Update path in DB
        async with get_db() as db:
            await db.execute(
                "UPDATE invoices SET pdf_path = ?, updated_at = CURRENT_TIMESTAMP WHERE id = ?",
                (pdf_path, invoice_id),
            )
            await db.commit()

    # Get SMTP config
    smtp_config = await _get_smtp_config()
    if not smtp_config.get("smtp_host") or not smtp_config.get("smtp_user"):
        raise HTTPException(
            status_code=400,
            detail="Configurare SMTP lipsa. Seteaza smtp_host, smtp_port, smtp_user, smtp_pass, smtp_from in setarile AI.",
        )

    # Build email
    subject = data.subject or f"Factura {inv_number} — CIP Inspection SRL"
    body = data.body or (
        f"Buna ziua,\n\n"
        f"Va transmitem atasat factura {inv_number} in valoare de {invoice['total']:.2f} RON.\n\n"
        f"Cu stima,\nCIP Inspection SRL"
    )

    msg = MIMEMultipart()
    msg["From"] = smtp_config.get("smtp_from") or smtp_config["smtp_user"]
    msg["To"] = data.to_email
    msg["Subject"] = subject
    msg.attach(MIMEText(body, "plain", "utf-8"))

    # Attach PDF
    pdf_file = Path(pdf_path)
    if pdf_file.exists():
        with open(pdf_file, "rb") as f:
            part = MIMEApplication(f.read(), Name=pdf_file.name)
            part["Content-Disposition"] = f'attachment; filename="{pdf_file.name}"'
            msg.attach(part)

    # Send email
    try:
        port = int(smtp_config.get("smtp_port") or "587")
        if port == 465:
            server = smtplib.SMTP_SSL(smtp_config["smtp_host"], port, timeout=30)
        else:
            server = smtplib.SMTP(smtp_config["smtp_host"], port, timeout=30)
            server.starttls()
        server.login(smtp_config["smtp_user"], smtp_config["smtp_pass"])
        server.sendmail(msg["From"], [data.to_email], msg.as_string())
        server.quit()
    except smtplib.SMTPException as exc:
        logger.error("Eroare SMTP: %s", exc)
        raise HTTPException(status_code=500, detail=f"Eroare la trimitere email: {exc}")
    except Exception as exc:
        logger.error("Eroare email: %s", exc)
        raise HTTPException(status_code=500, detail=f"Eroare la trimitere email: {exc}")

    await log_activity(
        action="invoice.send_email",
        summary=f"Factura {inv_number} trimisa pe email la {data.to_email}",
        details={
            "invoice_id": invoice_id,
            "invoice_number": inv_number,
            "to": data.to_email,
        },
    )

    return {
        "message": f"Factura {inv_number} trimisa cu succes la {data.to_email}.",
        "invoice_number": inv_number,
        "to": data.to_email,
    }


# ═══════════════════════════════════════════
# ANAF CUI Verification (free, no auth)
# ═══════════════════════════════════════════

@router.get("/verify-cui/{cui}")
async def verify_cui(cui: str):
    """Verificare CUI la ANAF — returnează date firmă (denumire, adresa, TVA, stare)."""
    # Clean CUI — remove RO prefix and spaces
    clean_cui = cui.strip().upper().replace("RO", "").strip()
    if not clean_cui.isdigit():
        raise HTTPException(400, "CUI invalid — trebuie să conțină doar cifre")

    today = date.today().strftime("%Y-%m-%d")
    payload = [{"cui": int(clean_cui), "data": today}]

    try:
        async with httpx.AsyncClient(timeout=10) as client:
            resp = await client.post(
                "https://webservicesp.anaf.ro/api/PlatitorTvaRest/api/v8/ws/tva",
                json=payload,
            )
            resp.raise_for_status()

        data = resp.json()
        found = data.get("found", [])
        if not found:
            return {"found": False, "cui": clean_cui, "message": "CUI negăsit în baza ANAF"}

        info = found[0]
        return {
            "found": True,
            "cui": clean_cui,
            "denumire": info.get("date_generale", {}).get("denumire", ""),
            "adresa": info.get("date_generale", {}).get("adresa", ""),
            "telefon": info.get("date_generale", {}).get("telefon", ""),
            "cod_postal": info.get("date_generale", {}).get("codPostal", ""),
            "stare": info.get("date_generale", {}).get("stare_inregistrare", ""),
            "tva": info.get("inregistrare_scop_Tva", {}).get("scpTVA", False),
            "data_verificare": today,
        }
    except httpx.HTTPError as exc:
        logger.error("Eroare ANAF API: %s", exc)
        raise HTTPException(502, f"Nu s-a putut verifica CUI la ANAF: {exc}")


