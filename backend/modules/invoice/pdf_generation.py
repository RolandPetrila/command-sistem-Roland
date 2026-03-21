"""
Invoice PDF Generation sub-router:
  - POST /api/invoice/{invoice_id}/pdf   — Generate invoice PDF
  - POST /api/invoice/offer-pdf          — Generate offer/quote PDF
  - GET  /api/invoice/export/csv         — Export invoices as CSV
  - GET  /api/invoice/export/excel       — Export invoices as Excel
  - GET  /api/invoice/templates          — List document templates
  - POST /api/invoice/templates/generate — Generate document from template
"""

from __future__ import annotations

import csv
import io
import json
import logging
from datetime import date
from pathlib import Path
from typing import Any

from fastapi import APIRouter, HTTPException
from fastapi.responses import FileResponse, StreamingResponse

from app.core.activity_log import log_activity
from app.db.database import get_db

from .router import (
    InvoiceItem,
    OfferPdfRequest,
    TemplateGenerateRequest,
    INVOICES_DIR,
    TEMPLATES_DIR,
    _build_invoice_pdf,
    _items_from_json,
)

logger = logging.getLogger(__name__)

pdf_router = APIRouter(prefix="/api/invoice", tags=["Invoice PDF"])


# ═══════════════════════════════════════════
# Export CSV
# ═══════════════════════════════════════════

@pdf_router.get("/export/csv")
async def export_csv():
    """Export toate facturile ca fisier CSV."""
    async with get_db() as db:
        cursor = await db.execute(
            """SELECT i.id, i.invoice_number, i.date, i.due_date, i.subtotal,
                      i.vat_percent, i.vat_amount, i.total, i.currency, i.status,
                      i.notes, i.created_at,
                      c.name as client_name, c.cui as client_cui
               FROM invoices i
               LEFT JOIN clients c ON i.client_id = c.id
               ORDER BY i.date DESC, i.id DESC"""
        )
        rows = await cursor.fetchall()

    output = io.StringIO()
    writer = csv.writer(output, delimiter=";", quoting=csv.QUOTE_MINIMAL)

    writer.writerow([
        "ID", "Nr. Factura", "Data", "Scadenta", "Client", "CUI Client",
        "Subtotal", "TVA %", "TVA Suma", "Total", "Moneda", "Status", "Observatii", "Creat la",
    ])

    for row in rows:
        writer.writerow([
            row["id"], row["invoice_number"], row["date"], row["due_date"] or "",
            row["client_name"] or "", row["client_cui"] or "",
            row["subtotal"], row["vat_percent"], row["vat_amount"],
            row["total"], row["currency"], row["status"],
            row["notes"] or "", row["created_at"],
        ])

    content = output.getvalue()
    output.close()

    bom = "\ufeff"
    buffer = io.BytesIO((bom + content).encode("utf-8"))

    await log_activity(
        action="invoice.export_csv",
        summary=f"Export CSV: {len(rows)} facturi",
        details={"count": len(rows)},
    )

    return StreamingResponse(
        buffer,
        media_type="text/csv; charset=utf-8",
        headers={
            "Content-Disposition": f'attachment; filename="facturi_export_{date.today().isoformat()}.csv"',
        },
    )


# ═══════════════════════════════════════════
# Export Excel
# ═══════════════════════════════════════════

@pdf_router.get("/export/excel")
async def export_excel():
    """Export toate facturile ca fisier Excel (.xlsx)."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    async with get_db() as db:
        cursor = await db.execute(
            """SELECT i.id, i.invoice_number, i.date, i.due_date, i.subtotal,
                      i.vat_percent, i.vat_amount, i.total, i.currency, i.status,
                      i.notes, i.created_at,
                      c.name as client_name, c.cui as client_cui
               FROM invoices i
               LEFT JOIN clients c ON i.client_id = c.id
               ORDER BY i.date DESC, i.id DESC"""
        )
        rows = await cursor.fetchall()

    wb = Workbook()
    ws = wb.active
    ws.title = "Facturi"

    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_fill = PatternFill(start_color="1a365d", end_color="1a365d", fill_type="solid")
    header_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    thin_border = Border(
        left=Side(style="thin", color="CBD5E0"),
        right=Side(style="thin", color="CBD5E0"),
        top=Side(style="thin", color="CBD5E0"),
        bottom=Side(style="thin", color="CBD5E0"),
    )

    headers = [
        "ID", "Nr. Factura", "Data", "Scadenta", "Client", "CUI Client",
        "Subtotal", "TVA %", "TVA Suma", "Total", "Moneda", "Status", "Observatii", "Creat la",
    ]

    for col_idx, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_idx, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align
        cell.border = thin_border

    for row_idx, row in enumerate(rows, 2):
        values = [
            row["id"], row["invoice_number"], row["date"], row["due_date"] or "",
            row["client_name"] or "", row["client_cui"] or "",
            row["subtotal"], row["vat_percent"], row["vat_amount"],
            row["total"], row["currency"], row["status"],
            row["notes"] or "", row["created_at"],
        ]
        for col_idx, val in enumerate(values, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=val)
            cell.border = thin_border

    col_widths = [6, 18, 12, 12, 25, 15, 12, 8, 12, 12, 8, 10, 30, 20]
    for i, width in enumerate(col_widths, 1):
        ws.column_dimensions[ws.cell(row=1, column=i).column_letter].width = width

    ws.auto_filter.ref = f"A1:N{len(rows) + 1}"

    total_row = len(rows) + 2
    ws.cell(row=total_row, column=5, value="TOTAL:").font = Font(bold=True, size=11)
    ws.cell(row=total_row, column=10, value=sum(r["total"] for r in rows)).font = Font(bold=True, size=11)

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)

    await log_activity(
        action="invoice.export_excel",
        summary=f"Export Excel: {len(rows)} facturi",
        details={"count": len(rows)},
    )

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={
            "Content-Disposition": f'attachment; filename="facturi_export_{date.today().isoformat()}.xlsx"',
        },
    )


# ═══════════════════════════════════════════
# Document templates
# ═══════════════════════════════════════════

@pdf_router.get("/templates")
async def list_templates():
    """Lista template-uri documente disponibile."""
    async with get_db() as db:
        cursor = await db.execute(
            "SELECT * FROM document_templates ORDER BY template_type, name"
        )
        rows = await cursor.fetchall()

    result = []
    for row in rows:
        entry = dict(row)
        try:
            entry["fields"] = json.loads(entry.pop("fields_json"))
        except (json.JSONDecodeError, KeyError):
            entry["fields"] = []
        result.append(entry)

    return result


@pdf_router.post("/templates/generate")
async def generate_template(data: TemplateGenerateRequest):
    """Genereaza document din template (contract, oferta, chitanta)."""
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    async with get_db() as db:
        cursor = await db.execute(
            "SELECT * FROM document_templates WHERE name = ?",
            (data.template_name,),
        )
        template = await cursor.fetchone()

    if not template:
        raise HTTPException(status_code=404, detail="Template negasit.")

    template = dict(template)
    template_type = template["template_type"]

    doc = DocxDocument()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    today_str = data.data.get("date", date.today().isoformat())

    if template_type == "contract":
        _build_contract_doc(doc, data.data, today_str)
    elif template_type == "offer":
        _build_offer_doc(doc, data.data, today_str)
    elif template_type == "receipt":
        _build_receipt_doc(doc, data.data, today_str)
    else:
        raise HTTPException(status_code=400, detail=f"Tip template necunoscut: {template_type}")

    filename = f"{data.template_name}_{today_str}"
    docx_path = TEMPLATES_DIR / f"{filename}.docx"
    doc.save(str(docx_path))

    await log_activity(
        action="invoice.template_generate",
        summary=f"Document generat: {template['display_name']}",
        details={
            "template": data.template_name,
            "type": template_type,
            "filename": f"{filename}.docx",
        },
    )

    return FileResponse(
        path=str(docx_path),
        filename=f"{filename}.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


# ═══════════════════════════════════════════
# Offer / Quote PDF
# ═══════════════════════════════════════════

@pdf_router.post("/offer-pdf")
async def generate_offer_pdf(data: OfferPdfRequest):
    """Generate an offer/quote PDF for CIP Inspection SRL."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm
        from reportlab.pdfgen import canvas as pdf_canvas
    except ImportError:
        return await _generate_offer_fitz(data)

    buf = io.BytesIO()
    c = pdf_canvas.Canvas(buf, pagesize=A4)
    width, height = A4

    c.setFont("Helvetica-Bold", 16)
    c.drawString(2 * cm, height - 2 * cm, "OFERTA DE PRET")
    c.setFont("Helvetica", 10)
    c.drawString(2 * cm, height - 2.8 * cm, "CIP Inspection SRL | CUI 43978110")
    c.drawString(2 * cm, height - 3.4 * cm, f"Data: {date.today().strftime('%d.%m.%Y')}")
    c.drawString(2 * cm, height - 3.9 * cm, f"Valabilitate: {data.validity_days} zile")

    c.setFont("Helvetica-Bold", 11)
    c.drawString(2 * cm, height - 5 * cm, f"Catre: {data.client_name}")
    if data.client_address:
        c.setFont("Helvetica", 10)
        c.drawString(2 * cm, height - 5.5 * cm, data.client_address)

    y = height - 7 * cm
    c.setFont("Helvetica-Bold", 10)
    c.drawString(2 * cm, y, "Nr.")
    c.drawString(3 * cm, y, "Descriere")
    c.drawString(12 * cm, y, "Cant.")
    c.drawString(14 * cm, y, "Pret unit.")
    c.drawString(17 * cm, y, "Total")
    y -= 0.6 * cm

    c.setFont("Helvetica", 10)
    subtotal = 0
    for i, item in enumerate(data.items, 1):
        total_item = round(item.quantity * item.unit_price, 2)
        subtotal += total_item
        c.drawString(2 * cm, y, str(i))
        c.drawString(3 * cm, y, item.description[:50])
        c.drawString(12 * cm, y, str(item.quantity))
        c.drawString(14 * cm, y, f"{item.unit_price:.2f}")
        c.drawString(17 * cm, y, f"{total_item:.2f}")
        y -= 0.5 * cm

    y -= 0.5 * cm
    c.setFont("Helvetica-Bold", 11)
    c.drawString(14 * cm, y, f"TOTAL: {subtotal:.2f} RON")

    if data.notes:
        y -= 1.5 * cm
        c.setFont("Helvetica", 9)
        c.drawString(2 * cm, y, f"Observatii: {data.notes}")

    c.save()
    buf.seek(0)

    await log_activity(
        action="invoice.offer_pdf",
        summary=f"Oferta PDF generata: {data.client_name} — {subtotal:.2f} RON",
        details={"client": data.client_name, "total": subtotal, "items": len(data.items)},
    )

    filename = f"oferta_{data.client_name.replace(' ', '_')}_{date.today().isoformat()}.pdf"
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


async def _generate_offer_fitz(data: OfferPdfRequest):
    """Fallback offer PDF using PyMuPDF (fitz) if reportlab not available."""
    import fitz as fitz_mod

    doc = fitz_mod.open()
    page = doc.new_page()

    y = 50
    page.insert_text((50, y), "OFERTA DE PRET", fontsize=18, fontname="helv")
    y += 25
    page.insert_text((50, y), "CIP Inspection SRL | CUI 43978110", fontsize=10)
    y += 15
    page.insert_text((50, y), f"Data: {date.today().strftime('%d.%m.%Y')} | Valabilitate: {data.validity_days} zile", fontsize=10)
    y += 25
    page.insert_text((50, y), f"Catre: {data.client_name}", fontsize=12, fontname="helv")
    if data.client_address:
        y += 15
        page.insert_text((50, y), data.client_address, fontsize=10)
    y += 30

    subtotal = 0
    for i, item in enumerate(data.items, 1):
        total_item = round(item.quantity * item.unit_price, 2)
        subtotal += total_item
        line = f"{i}. {item.description} — {item.quantity} x {item.unit_price:.2f} = {total_item:.2f} RON"
        page.insert_text((50, y), line, fontsize=10)
        y += 15

    y += 10
    page.insert_text((50, y), f"TOTAL: {subtotal:.2f} RON", fontsize=13, fontname="helv")

    if data.notes:
        y += 25
        page.insert_text((50, y), f"Observatii: {data.notes}", fontsize=9)

    buf = io.BytesIO()
    doc.save(buf)
    doc.close()
    buf.seek(0)

    await log_activity(
        action="invoice.offer_pdf",
        summary=f"Oferta PDF (fitz): {data.client_name} — {subtotal:.2f} RON",
    )

    filename = f"oferta_{data.client_name.replace(' ', '_')}_{date.today().isoformat()}.pdf"
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ═══════════════════════════════════════════
# Invoice PDF generation endpoint
# ═══════════════════════════════════════════

@pdf_router.post("/{invoice_id}/pdf")
async def generate_pdf(invoice_id: int):
    """Genereaza PDF profesional si returneaza fisierul."""
    async with get_db() as db:
        cursor = await db.execute(
            """SELECT i.*, c.name as client_name, c.cui as client_cui,
                      c.reg_com as client_reg_com, c.address as client_address,
                      c.email as client_email, c.phone as client_phone
               FROM invoices i
               LEFT JOIN clients c ON i.client_id = c.id
               WHERE i.id = ?""",
            (invoice_id,),
        )
        row = await cursor.fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Factura negasita.")

    invoice = dict(row)
    items = _items_from_json(invoice.get("items_json", "[]"))

    pdf_filename = f"{invoice['invoice_number']}.pdf"
    pdf_path = INVOICES_DIR / pdf_filename

    try:
        _build_invoice_pdf(invoice, items, str(pdf_path))
    except Exception as exc:
        logger.error("Eroare la generare PDF: %s", exc)
        raise HTTPException(status_code=500, detail=f"Eroare la generare PDF: {exc}")

    async with get_db() as db:
        await db.execute(
            "UPDATE invoices SET pdf_path = ?, updated_at = CURRENT_TIMESTAMP WHERE id = ?",
            (str(pdf_path), invoice_id),
        )
        await db.commit()

    await log_activity(
        action="invoice.pdf",
        summary=f"PDF generat: {invoice['invoice_number']}",
        details={"invoice_id": invoice_id, "pdf_path": str(pdf_path)},
    )
    return FileResponse(
        path=str(pdf_path),
        filename=pdf_filename,
        media_type="application/pdf",
    )


# ═══════════════════════════════════════════
# Template builder helpers (docx)
# These are called from generate_template above.
# ═══════════════════════════════════════════

def _build_contract_doc(doc, data: dict, today_str: str) -> None:
    """Build a contract DOCX document."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    heading = doc.add_heading("CONTRACT DE PRESTARI SERVICII", 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Data: {today_str}")
    doc.add_paragraph("")

    doc.add_heading("Parti contractante", level=1)
    prestator = data.get("prestator", "CIP Inspection SRL, CUI 43978110")
    beneficiar = data.get("beneficiar", "[Beneficiar]")
    doc.add_paragraph(f"Prestator: {prestator}")
    doc.add_paragraph(f"Beneficiar: {beneficiar}")

    doc.add_heading("Obiectul contractului", level=1)
    obiect = data.get("obiect", "Servicii de traducere tehnica autorizata")
    doc.add_paragraph(obiect)

    doc.add_heading("Valoare si modalitate de plata", level=1)
    valoare = data.get("valoare", "[Valoare]")
    doc.add_paragraph(f"Valoare: {valoare} RON")

    doc.add_heading("Termen de executie", level=1)
    termen = data.get("termen", "[Termen]")
    doc.add_paragraph(termen)

    doc.add_paragraph("")
    doc.add_paragraph("Prezentul contract a fost incheiat in 2 exemplare originale.")
    doc.add_paragraph("")
    semnatura_table = doc.add_table(rows=2, cols=2)
    semnatura_table.cell(0, 0).text = "Prestator"
    semnatura_table.cell(0, 1).text = "Beneficiar"
    semnatura_table.cell(1, 0).text = "Semnatura: _______________"
    semnatura_table.cell(1, 1).text = "Semnatura: _______________"


def _build_offer_doc(doc, data: dict, today_str: str) -> None:
    """Build an offer DOCX document."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    heading = doc.add_heading("OFERTA DE PRET", 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Data: {today_str}")
    validity = data.get("validity_days", 30)
    doc.add_paragraph(f"Valabilitate: {validity} zile")
    doc.add_paragraph("")

    client = data.get("client_name", "[Client]")
    doc.add_heading(f"Catre: {client}", level=1)

    doc.add_heading("Servicii oferite", level=1)
    items = data.get("items", [])
    if items:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Descriere"
        hdr[1].text = "Cantitate"
        hdr[2].text = "Pret unitar (RON)"
        hdr[3].text = "Total (RON)"
        subtotal = 0
        for item in items:
            row = table.add_row().cells
            row[0].text = str(item.get("description", ""))
            row[1].text = str(item.get("quantity", 1))
            row[2].text = f"{item.get('unit_price', 0):.2f}"
            t = round(item.get("quantity", 1) * item.get("unit_price", 0), 2)
            row[3].text = f"{t:.2f}"
            subtotal += t
        doc.add_paragraph(f"\nTotal: {subtotal:.2f} RON")

    notes = data.get("notes", "")
    if notes:
        doc.add_heading("Observatii", level=2)
        doc.add_paragraph(notes)


def _build_receipt_doc(doc, data: dict, today_str: str) -> None:
    """Build a receipt DOCX document."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    heading = doc.add_heading("CHITANTA", 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Nr: {data.get('nr', '[Nr]')}  |  Data: {today_str}")
    doc.add_paragraph("")
    doc.add_paragraph(f"Am primit de la: {data.get('de_la', '[Platitor]')}")
    doc.add_paragraph(f"Suma de: {data.get('suma', '[Suma]')} RON")
    doc.add_paragraph(f"Reprezentand: {data.get('reprezentand', '[Descriere]')}")
    doc.add_paragraph("")
    doc.add_paragraph("Semnatura: _______________")
