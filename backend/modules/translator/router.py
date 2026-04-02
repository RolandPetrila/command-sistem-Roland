"""
Translator API endpoints — traducere text/fisiere, TM, glosar, detectie limba, quality check.

Endpoints:
  POST   /api/translator/text            — Traducere text (cu cache)
  POST   /api/translator/file            — Traducere fisier (PDF/DOCX/TXT)
  POST   /api/translator/detect          — Detectie limba
  POST   /api/translator/quality-check   — Evaluare calitate traducere (AI)
  POST   /api/translator/compare         — Comparatie 2 provideri
  GET    /api/translator/tm              — Lista intrari TM
  POST   /api/translator/tm              — Adaugare intrare TM
  GET    /api/translator/tm/search       — Cautare in TM
  GET    /api/translator/tm/stats        — Statistici TM
  DELETE /api/translator/tm/{id}         — Sterge intrare TM
  GET    /api/translator/glossary        — Lista termeni glosar
  POST   /api/translator/glossary        — Adaugare termen
  PUT    /api/translator/glossary/{id}   — Actualizare termen
  DELETE /api/translator/glossary/{id}   — Stergere termen
  POST   /api/translator/glossary/import — Import CSV
  GET    /api/translator/glossary/export — Export CSV glosar
  GET    /api/translator/usage           — Statistici utilizare provideri
  GET    /api/translator/providers       — Lista provideri disponibili
  GET    /api/translator/history         — Istoric traduceri cu search/filtrare
"""

from __future__ import annotations

import asyncio
import csv
import io
import logging
import os
import time
from io import BytesIO
from pathlib import Path

# fitz, docx — lazy imported inside functions to cut cold-start time
from fastapi import APIRouter, File, Form, HTTPException, Query, UploadFile
from fastapi.responses import StreamingResponse

from app.core.activity_log import log_activity
from app.db.database import get_db

from .models import (
    CompareRequest,
    DetectRequest,
    GlossaryAddRequest,
    GlossaryUpdateRequest,
    QualityCheckRequest,
    TMAddRequest,
    TranslateTextRequest,
)
from .helpers import (
    compute_cache_hash,
    ensure_cache_table,
    extract_text_from_file,
    get_provider_latencies,
    record_latency,
    save_upload,
)
from .providers import (
    detect_language,
    get_available_translation_providers,
    get_deepl_usage,
    translate_with_chain,
)
from .tm import (
    add_to_tm,
    delete_tm_entry,
    get_tm_entries,
    get_tm_stats,
    search_tm,
    translate_with_tm,
)
from .glossary import (
    add_term,
    apply_glossary,
    delete_term,
    get_glossary,
    get_glossary_domains,
    import_glossary_csv,
    update_term,
)

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/api/translator", tags=["Traducator"])


# ---------------------------------------------------------------------------
# POST /api/translator/text — Translate text
# ---------------------------------------------------------------------------

@router.post("/text")
async def translate_text_endpoint(req: TranslateTextRequest):
    """Traducere text cu suport TM, glosar si cache."""
    if not req.text.strip():
        raise HTTPException(400, "Textul nu poate fi gol")

    chars_count = len(req.text)
    from_cache = False

    # Step 0: Check translation cache
    await ensure_cache_table()
    cache_hash = compute_cache_hash(req.text, req.source_lang, req.target_lang)

    async with get_db() as db:
        cursor = await db.execute(
            "SELECT target_text, provider FROM translation_cache WHERE hash = ?",
            (cache_hash,),
        )
        cached = await cursor.fetchone()

    if cached:
        translated_text = cached["target_text"]
        provider = cached["provider"] + " (cache)"
        from_cache = True
        tm_hits = 0
        glossary_replacements = []

        await log_activity(
            action="translator.text",
            summary=f"Traducere din cache {req.source_lang.upper()}->{req.target_lang.upper()}: {chars_count} car.",
            details={"from_cache": True, "chars_count": chars_count},
        )

        return {
            "translated_text": translated_text,
            "provider": provider,
            "source_lang": req.source_lang,
            "target_lang": req.target_lang,
            "chars_count": chars_count,
            "tm_hits": 0,
            "glossary_replacements": [],
            "from_cache": True,
        }

    # Step 1: Apply glossary if enabled
    text_to_translate = req.text
    glossary_replacements = []
    tm_hits = 0

    if req.use_glossary:
        glossary_result = await apply_glossary(text_to_translate, req.source_lang, req.target_lang)
        if glossary_result["replacements"]:
            text_to_translate = glossary_result["text"]
            glossary_replacements = glossary_result["replacements"]

    # Step 2: Translate (with or without TM)
    provider_latency_ms = None
    if req.use_tm:
        result = await translate_with_tm(
            text=text_to_translate,
            source_lang=req.source_lang,
            target_lang=req.target_lang,
            translate_fn=lambda t, s, tl: translate_with_chain(t, s, tl, req.provider),
            domain=req.domain,
        )
        translated_text = result["translated_text"]
        provider = result["provider"]
        tm_hits = result["tm_hits"]
    else:
        try:
            result = await translate_with_chain(
                text_to_translate, req.source_lang, req.target_lang, req.provider
            )
            translated_text = result["translated_text"]
            provider = result["provider"]
            provider_latency_ms = result.get("provider_latency_ms")
        except RuntimeError as e:
            raise HTTPException(503, str(e))

    # R4-14: Record provider latency
    if provider_latency_ms is not None:
        record_latency(provider, provider_latency_ms)

    # Step 3: Store in cache
    async with get_db() as db:
        await db.execute(
            """
            INSERT OR REPLACE INTO translation_cache (hash, source_text, target_text, source_lang, target_lang, provider)
            VALUES (?, ?, ?, ?, ?, ?)
            """,
            (cache_hash, req.text, translated_text, req.source_lang, req.target_lang, provider),
        )
        await db.commit()

    # Step 4: Log translation in history
    async with get_db() as db:
        await db.execute(
            """
            INSERT INTO translations (source_text, target_text, source_lang, target_lang, provider, chars_count)
            VALUES (?, ?, ?, ?, ?, ?)
            """,
            (req.text, translated_text, req.source_lang, req.target_lang, provider, chars_count),
        )
        await db.commit()

    # Step 5: Auto-populate TM if enabled
    if req.auto_tm and translated_text and len(req.text.strip()) >= 5:
        try:
            await add_to_tm(
                source=req.text.strip(),
                target=translated_text.strip(),
                source_lang=req.source_lang,
                target_lang=req.target_lang,
                domain=req.domain,
            )
            logger.debug("Auto-TM: segment adaugat (%s->%s)", req.source_lang, req.target_lang)
        except Exception as e:
            logger.warning("Auto-TM: eroare la adaugare segment: %s", e)

    await log_activity(
        action="translator.text",
        summary=f"Traducere {req.source_lang.upper()}->{req.target_lang.upper()}: {chars_count} caractere ({provider})",
        details={
            "source_lang": req.source_lang,
            "target_lang": req.target_lang,
            "provider": provider,
            "chars_count": chars_count,
            "tm_hits": tm_hits,
            "glossary_replacements": len(glossary_replacements),
            "auto_tm": req.auto_tm,
            "from_cache": False,
        },
    )

    return {
        "translated_text": translated_text,
        "provider": provider,
        "source_lang": req.source_lang,
        "target_lang": req.target_lang,
        "chars_count": chars_count,
        "tm_hits": tm_hits,
        "glossary_replacements": glossary_replacements,
        "from_cache": False,
        "provider_latency_ms": provider_latency_ms,
    }


# ---------------------------------------------------------------------------
# POST /api/translator/file — Translate file
# ---------------------------------------------------------------------------

@router.post("/file")
async def translate_file_endpoint(
    file: UploadFile = File(...),
    source_lang: str = Form("en"),
    target_lang: str = Form("ro"),
    provider: str = Form("auto"),
    use_tm: bool = Form(True),
    use_glossary: bool = Form(True),
):
    """
    Traducere fisier (PDF/DOCX/TXT).

    - DOCX: traduce paragrafele in-place, returneaza DOCX tradus
    - PDF: extrage text, traduce, returneaza DOCX (reconstituirea layout-ului PDF e prea complexa)
    - TXT: traduce continutul, returneaza TXT
    """
    if not file.filename:
        raise HTTPException(400, "Fisierul nu are nume")

    ext = Path(file.filename).suffix.lower()
    if ext not in (".pdf", ".docx", ".txt", ".md"):
        raise HTTPException(
            400,
            f"Format nesuportat: {ext}. Acceptate: PDF, DOCX, TXT, MD"
        )

    tmp_path = await save_upload(file)

    try:
        total_chars = 0
        tm_hits_total = 0

        if ext == ".docx":
            # Translate DOCX paragraphs in-place
            from docx import Document as DocxDocument
            doc = DocxDocument(tmp_path)

            # Helper: translate a single text through TM/glossary/chain
            async def _translate_one(text_in: str) -> str:
                nonlocal total_chars, tm_hits_total
                if not text_in.strip():
                    return text_in
                total_chars += len(text_in)
                t = text_in
                if use_glossary:
                    g_result = await apply_glossary(t, source_lang, target_lang)
                    t = g_result["text"]
                if use_tm:
                    result = await translate_with_tm(
                        text=t, source_lang=source_lang, target_lang=target_lang,
                        translate_fn=lambda tx, s, tl: translate_with_chain(tx, s, tl, provider),
                    )
                    tm_hits_total += result["tm_hits"]
                    return result["translated_text"]
                result = await translate_with_chain(t, source_lang, target_lang, provider)
                return result["translated_text"]

            # Translate paragraphs
            for para in doc.paragraphs:
                if not para.text.strip():
                    continue
                translated = await _translate_one(para.text)
                if para.runs:
                    para.runs[0].text = translated
                    for run in para.runs[1:]:
                        run.text = ""
                else:
                    para.text = translated

            # C3: Translate tables (cell by cell, preserving structure)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if not cell.text.strip():
                            continue
                        translated = await _translate_one(cell.text)
                        # Preserve cell formatting: set first paragraph, clear rest
                        if cell.paragraphs:
                            first_para = cell.paragraphs[0]
                            if first_para.runs:
                                first_para.runs[0].text = translated
                                for run in first_para.runs[1:]:
                                    run.text = ""
                            else:
                                first_para.text = translated
                            for extra_para in cell.paragraphs[1:]:
                                extra_para.text = ""

            # Save translated DOCX to buffer
            output_buffer = BytesIO()
            doc.save(output_buffer)
            output_buffer.seek(0)

            output_filename = Path(file.filename).stem + f"_{target_lang.upper()}.docx"

            await log_activity(
                action="translator.file",
                summary=f"Fisier tradus: {file.filename} ({total_chars} car.)",
                details={
                    "filename": file.filename,
                    "source_lang": source_lang,
                    "target_lang": target_lang,
                    "chars_count": total_chars,
                    "tm_hits": tm_hits_total,
                },
            )

            return StreamingResponse(
                output_buffer,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f'attachment; filename="{output_filename}"'},
            )

        elif ext == ".pdf":
            # PDF: extract text, translate in batches, return as DOCX
            import fitz  # PyMuPDF
            from docx import Document as DocxDocument
            BATCH_SIZE = 10  # paragraphs per API call

            pdf_doc = fitz.open(tmp_path)
            output_doc = DocxDocument()

            for page_num, page in enumerate(pdf_doc):
                if page_num > 0:
                    output_doc.add_page_break()

                text = page.get_text()
                if not text.strip():
                    continue

                paragraphs = [p for p in text.split("\n") if p.strip()]

                # C2: Batch paragraphs — translate groups of BATCH_SIZE at once
                for batch_start in range(0, len(paragraphs), BATCH_SIZE):
                    batch = paragraphs[batch_start:batch_start + BATCH_SIZE]
                    # Join batch with separator for single API call
                    separator = "\n\n---PARA---\n\n"
                    combined_text = separator.join(batch)
                    total_chars += sum(len(p) for p in batch)

                    text_to_translate = combined_text

                    if use_glossary:
                        g_result = await apply_glossary(text_to_translate, source_lang, target_lang)
                        text_to_translate = g_result["text"]

                    if use_tm:
                        result = await translate_with_tm(
                            text=text_to_translate,
                            source_lang=source_lang,
                            target_lang=target_lang,
                            translate_fn=lambda t, s, tl: translate_with_chain(t, s, tl, provider),
                        )
                        translated_combined = result["translated_text"]
                        tm_hits_total += result["tm_hits"]
                    else:
                        result = await translate_with_chain(text_to_translate, source_lang, target_lang, provider)
                        translated_combined = result["translated_text"]

                    # Split back into individual paragraphs
                    translated_parts = translated_combined.split("---PARA---")
                    for i, part in enumerate(translated_parts):
                        output_doc.add_paragraph(part.strip())

            pdf_doc.close()

            output_buffer = BytesIO()
            output_doc.save(output_buffer)
            output_buffer.seek(0)

            output_filename = Path(file.filename).stem + f"_{target_lang.upper()}.docx"

            await log_activity(
                action="translator.file",
                summary=f"PDF tradus: {file.filename} ({total_chars} car.) -> DOCX",
                details={
                    "filename": file.filename,
                    "source_lang": source_lang,
                    "target_lang": target_lang,
                    "chars_count": total_chars,
                    "tm_hits": tm_hits_total,
                    "output_format": "docx",
                },
            )

            return StreamingResponse(
                output_buffer,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f'attachment; filename="{output_filename}"'},
            )

        else:
            # TXT / MD: translate and return as TXT
            with open(tmp_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()

            total_chars = len(content)
            text_to_translate = content

            if use_glossary:
                g_result = await apply_glossary(text_to_translate, source_lang, target_lang)
                text_to_translate = g_result["text"]

            if use_tm:
                result = await translate_with_tm(
                    text=text_to_translate,
                    source_lang=source_lang,
                    target_lang=target_lang,
                    translate_fn=lambda t, s, tl: translate_with_chain(t, s, tl, provider),
                )
                translated = result["translated_text"]
                tm_hits_total = result["tm_hits"]
            else:
                result = await translate_with_chain(text_to_translate, source_lang, target_lang, provider)
                translated = result["translated_text"]

            output_buffer = BytesIO(translated.encode("utf-8"))
            output_filename = Path(file.filename).stem + f"_{target_lang.upper()}.txt"

            await log_activity(
                action="translator.file",
                summary=f"TXT tradus: {file.filename} ({total_chars} car.)",
                details={
                    "filename": file.filename,
                    "source_lang": source_lang,
                    "target_lang": target_lang,
                    "chars_count": total_chars,
                    "tm_hits": tm_hits_total,
                },
            )

            return StreamingResponse(
                output_buffer,
                media_type="text/plain; charset=utf-8",
                headers={"Content-Disposition": f'attachment; filename="{output_filename}"'},
            )

    except HTTPException:
        raise
    except RuntimeError as e:
        raise HTTPException(503, str(e))
    except Exception as e:
        logger.error("Eroare la traducerea fisierului: %s", e, exc_info=True)
        raise HTTPException(500, f"Eroare la traducerea fisierului: {str(e)}")
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


# ---------------------------------------------------------------------------
# POST /api/translator/detect — Detect language
# ---------------------------------------------------------------------------

@router.post("/detect")
async def detect_language_endpoint(req: DetectRequest):
    """Detecteaza limba textului."""
    if not req.text.strip():
        raise HTTPException(400, "Textul nu poate fi gol")

    result = await detect_language(req.text)

    await log_activity(
        action="translator.detect",
        summary=f"Detectie limba: {result['language']} ({result['confidence']:.0%})",
    )

    return result


# ---------------------------------------------------------------------------
# GET /api/translator/tm — List TM entries
# ---------------------------------------------------------------------------

@router.get("/tm")
async def list_tm(
    source_lang: str | None = Query(None),
    target_lang: str | None = Query(None),
    domain: str | None = Query(None),
    skip: int = Query(0, ge=0),
    limit: int = Query(50, ge=1, le=500),
):
    """Lista intrari din Translation Memory cu paginare."""
    return await get_tm_entries(
        source_lang=source_lang,
        target_lang=target_lang,
        domain=domain,
        skip=skip,
        limit=limit,
    )


# ---------------------------------------------------------------------------
# POST /api/translator/tm — Add TM entry
# ---------------------------------------------------------------------------

@router.post("/tm")
async def add_tm_entry(req: TMAddRequest):
    """Adauga manual o intrare in Translation Memory."""
    if not req.source.strip() or not req.target.strip():
        raise HTTPException(400, "Sursa si traducerea nu pot fi goale")

    tm_id = await add_to_tm(
        source=req.source,
        target=req.target,
        source_lang=req.source_lang,
        target_lang=req.target_lang,
        domain=req.domain,
    )

    await log_activity(
        action="translator.tm_add",
        summary=f"TM: adaugat segment {req.source_lang.upper()}->{req.target_lang.upper()}",
        details={"tm_id": tm_id, "source_preview": req.source[:50]},
    )

    return {"id": tm_id, "status": "adaugat"}


# ---------------------------------------------------------------------------
# GET /api/translator/tm/search — Search TM
# ---------------------------------------------------------------------------

@router.get("/tm/search")
async def search_tm_endpoint(
    q: str = Query(..., min_length=2),
    source_lang: str = Query("en"),
    target_lang: str = Query("ro"),
    threshold: float = Query(0.7, ge=0.0, le=1.0),
    limit: int = Query(10, ge=1, le=100),
):
    """Cautare in Translation Memory dupa text similar. Returneaza scor similaritate per rezultat."""
    results = await search_tm(
        text=q,
        source_lang=source_lang,
        target_lang=target_lang,
        threshold=threshold,
        limit=limit,
    )
    # R4-12: Ensure each result has a normalized score field (0.0-1.0)
    for r in results:
        r["score"] = r.get("confidence", 0.0)
    return {"query": q, "results": results, "count": len(results)}


# ---------------------------------------------------------------------------
# GET /api/translator/tm/stats — TM statistics
# ---------------------------------------------------------------------------

@router.get("/tm/stats")
async def tm_stats():
    """Statistici Translation Memory."""
    return await get_tm_stats()


# ---------------------------------------------------------------------------
# DELETE /api/translator/tm/{tm_id} — Delete TM entry
# ---------------------------------------------------------------------------

@router.delete("/tm/{tm_id}")
async def delete_tm(tm_id: int):
    """Sterge o intrare din Translation Memory."""
    deleted = await delete_tm_entry(tm_id)
    if not deleted:
        raise HTTPException(404, "Intrarea TM nu a fost gasita")

    await log_activity(
        action="translator.tm_delete",
        summary=f"TM: segment #{tm_id} sters",
    )
    return {"status": "sters", "id": tm_id}


# ---------------------------------------------------------------------------
# GET /api/translator/glossary — List glossary terms
# ---------------------------------------------------------------------------

@router.get("/glossary")
async def list_glossary(
    domain: str | None = Query(None),
    search: str | None = Query(None),
    source_lang: str | None = Query(None),
    target_lang: str | None = Query(None),
    client_id: int | None = Query(None),
    skip: int = Query(0, ge=0),
    limit: int = Query(100, ge=1, le=1000),
):
    """Lista termeni glosar cu filtre optionale, inclusiv per client."""
    return await get_glossary(
        domain=domain,
        search=search,
        source_lang=source_lang,
        target_lang=target_lang,
        client_id=client_id,
        skip=skip,
        limit=limit,
    )


# ---------------------------------------------------------------------------
# POST /api/translator/glossary — Add term
# ---------------------------------------------------------------------------

@router.post("/glossary")
async def add_glossary_term(req: GlossaryAddRequest):
    """Adauga un termen in glosar."""
    if not req.source.strip() or not req.target.strip():
        raise HTTPException(400, "Sursa si traducerea nu pot fi goale")

    try:
        term = await add_term(
            source=req.source,
            target=req.target,
            source_lang=req.source_lang,
            target_lang=req.target_lang,
            domain=req.domain,
            notes=req.notes,
            client_id=req.client_id,
        )
    except ValueError as e:
        raise HTTPException(409, str(e))

    # Invalidate cache entries for the affected language pair
    await ensure_cache_table()
    async with get_db() as db:
        await db.execute(
            "DELETE FROM translation_cache WHERE source_lang = ? AND target_lang = ?",
            (req.source_lang.lower(), req.target_lang.lower()),
        )
        await db.commit()

    await log_activity(
        action="translator.glossary_add",
        summary=f"Glosar: '{req.source}' -> '{req.target}' ({req.domain})",
    )

    return term


# ---------------------------------------------------------------------------
# PUT /api/translator/glossary/{term_id} — Update term
# ---------------------------------------------------------------------------

@router.put("/glossary/{term_id}")
async def update_glossary_term(term_id: int, req: GlossaryUpdateRequest):
    """Actualizeaza un termen din glosar."""
    updated = await update_term(
        term_id=term_id,
        source=req.source,
        target=req.target,
        source_lang=req.source_lang,
        target_lang=req.target_lang,
        domain=req.domain,
        notes=req.notes,
    )
    if updated is None:
        raise HTTPException(404, "Termenul nu a fost gasit in glosar")

    # Invalidate cache entries for the affected language pair
    await ensure_cache_table()
    async with get_db() as db:
        await db.execute(
            "DELETE FROM translation_cache WHERE source_lang = ? AND target_lang = ?",
            (req.source_lang.lower(), req.target_lang.lower()),
        )
        await db.commit()

    await log_activity(
        action="translator.glossary_update",
        summary=f"Glosar: termen #{term_id} actualizat",
    )

    return updated


# ---------------------------------------------------------------------------
# DELETE /api/translator/glossary/{term_id} — Delete term
# ---------------------------------------------------------------------------

@router.delete("/glossary/{term_id}")
async def delete_glossary_term(term_id: int):
    """Sterge un termen din glosar."""
    deleted = await delete_term(term_id)
    if not deleted:
        raise HTTPException(404, "Termenul nu a fost gasit in glosar")

    await log_activity(
        action="translator.glossary_delete",
        summary=f"Glosar: termen #{term_id} sters",
    )

    return {"status": "sters", "id": term_id}


# ---------------------------------------------------------------------------
# POST /api/translator/glossary/import — Import CSV
# ---------------------------------------------------------------------------

@router.post("/glossary/import")
async def import_glossary(
    file: UploadFile = File(...),
    source_lang: str = Form("en"),
    target_lang: str = Form("ro"),
    domain: str = Form("general"),
):
    """Import termeni glosar din fisier CSV."""
    if not file.filename or not file.filename.lower().endswith(".csv"):
        raise HTTPException(400, "Doar fisiere CSV sunt acceptate")

    content = await file.read()
    try:
        text = content.decode("utf-8")
    except UnicodeDecodeError:
        try:
            text = content.decode("cp1252")
        except UnicodeDecodeError:
            raise HTTPException(400, "Nu s-a putut decodifica fisierul. Foloseste codificarea UTF-8.")

    result = await import_glossary_csv(
        file_content=text,
        source_lang=source_lang,
        target_lang=target_lang,
        default_domain=domain,
    )

    # Invalidate cache entries for the affected language pair
    await ensure_cache_table()
    async with get_db() as db:
        await db.execute(
            "DELETE FROM translation_cache WHERE source_lang = ? AND target_lang = ?",
            (source_lang.lower(), target_lang.lower()),
        )
        await db.commit()

    await log_activity(
        action="translator.glossary_import",
        summary=f"Import glosar: {result['imported']} termeni importati, {result['skipped']} omisi",
        details={
            "filename": file.filename,
            "imported": result["imported"],
            "skipped": result["skipped"],
        },
    )

    return result


# ---------------------------------------------------------------------------
# GET /api/translator/glossary/export — Export glossary as CSV
# ---------------------------------------------------------------------------

@router.get("/glossary/export")
async def export_glossary_csv():
    """Export intreg glosarul ca fisier CSV."""
    async with get_db() as db:
        cursor = await db.execute(
            """
            SELECT term_source, term_target, source_lang, target_lang, domain, notes
            FROM glossary_terms
            ORDER BY domain, term_source
            """
        )
        rows = await cursor.fetchall()

    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["source", "target", "source_lang", "target_lang", "domain", "notes"])

    for row in rows:
        writer.writerow([
            row["term_source"],
            row["term_target"],
            row["source_lang"],
            row["target_lang"],
            row["domain"],
            row["notes"] or "",
        ])

    output.seek(0)

    await log_activity(
        action="translator.glossary_export",
        summary=f"Export glosar CSV: {len(rows)} termeni",
    )

    return StreamingResponse(
        iter([output.getvalue()]),
        media_type="text/csv; charset=utf-8",
        headers={"Content-Disposition": 'attachment; filename="glossary_export.csv"'},
    )


# ---------------------------------------------------------------------------
# GET /api/translator/glossary/domains — List domains
# ---------------------------------------------------------------------------

@router.get("/glossary/domains")
async def list_glossary_domains():
    """Lista domeniilor din glosar."""
    domains = await get_glossary_domains()
    return {"domains": domains}


# ---------------------------------------------------------------------------
# GET /api/translator/usage — Provider usage stats
# ---------------------------------------------------------------------------

@router.get("/usage")
async def get_usage_stats():
    """Statistici utilizare provideri traducere (DeepL chars, etc.)."""
    stats = {}

    # DeepL usage
    deepl_usage = await get_deepl_usage()
    if deepl_usage:
        stats["deepl"] = deepl_usage

    # Translation history stats from DB
    async with get_db() as db:
        cursor = await db.execute(
            """
            SELECT provider, COUNT(*) as translations_count,
                   SUM(chars_count) as total_chars
            FROM translations
            GROUP BY provider
            ORDER BY total_chars DESC
            """
        )
        provider_stats = [dict(r) for r in await cursor.fetchall()]

        cursor = await db.execute(
            "SELECT COUNT(*) as cnt, SUM(chars_count) as total FROM translations"
        )
        totals = dict(await cursor.fetchone())

    stats["history"] = {
        "total_translations": totals["cnt"] or 0,
        "total_chars": totals["total"] or 0,
        "per_provider": provider_stats,
    }

    return stats


# ---------------------------------------------------------------------------
# GET /api/translator/providers — List available providers
# ---------------------------------------------------------------------------

@router.get("/providers")
async def list_translation_providers():
    """Lista provideri de traducere disponibili."""
    return await get_available_translation_providers()


# ---------------------------------------------------------------------------
# GET /api/translator/history — Translation history with search + filters
# ---------------------------------------------------------------------------

@router.get("/history")
async def translation_history(
    q: str = Query("", description="Cauta in text sursa/tradus"),
    provider: str = Query("", description="Filtreaza dupa provider"),
    source_lang: str = Query("", description="Filtreaza dupa limba sursa"),
    target_lang: str = Query("", description="Filtreaza dupa limba tinta"),
    date_from: str = Query("", description="De la data (YYYY-MM-DD)"),
    date_to: str = Query("", description="Pana la data (YYYY-MM-DD)"),
    page: int = Query(1, ge=1, description="Numar pagina"),
    per_page: int = Query(20, ge=1, le=100, description="Rezultate per pagina"),
):
    """Istoric traduceri cu cautare, filtre si paginare."""
    conditions: list[str] = []
    params: list = []

    if q.strip():
        conditions.append("(source_text LIKE ? OR target_text LIKE ?)")
        like_q = f"%{q.strip()}%"
        params.extend([like_q, like_q])

    if provider.strip():
        conditions.append("provider LIKE ?")
        params.append(f"%{provider.strip()}%")

    if source_lang.strip():
        conditions.append("source_lang = ?")
        params.append(source_lang.strip().lower())

    if target_lang.strip():
        conditions.append("target_lang = ?")
        params.append(target_lang.strip().lower())

    if date_from.strip():
        conditions.append("created_at >= ?")
        params.append(date_from.strip())

    if date_to.strip():
        conditions.append("created_at <= ?")
        params.append(date_to.strip() + " 23:59:59")

    where_clause = (" WHERE " + " AND ".join(conditions)) if conditions else ""
    offset = (page - 1) * per_page

    async with get_db() as db:
        cursor = await db.execute(
            f"""
            SELECT id, source_text, target_text, source_lang, target_lang,
                   provider, chars_count, file_name, created_at
            FROM translations
            {where_clause}
            ORDER BY created_at DESC
            LIMIT ? OFFSET ?
            """,
            (*params, per_page, offset),
        )
        rows = [dict(r) for r in await cursor.fetchall()]

        cursor_count = await db.execute(
            f"SELECT COUNT(*) as cnt FROM translations {where_clause}",
            params,
        )
        total = (await cursor_count.fetchone())["cnt"]

    # Truncate long texts for list view
    for row in rows:
        if row.get("source_text") and len(row["source_text"]) > 200:
            row["source_text"] = row["source_text"][:200] + "..."
        if row.get("target_text") and len(row["target_text"]) > 200:
            row["target_text"] = row["target_text"][:200] + "..."

    total_pages = max(1, (total + per_page - 1) // per_page)

    return {
        "items": rows,
        "total": total,
        "page": page,
        "per_page": per_page,
        "total_pages": total_pages,
    }


# ---------------------------------------------------------------------------
# R4-13: POST /api/translator/translate-batch — Batch multi-file translation
# ---------------------------------------------------------------------------

@router.post("/translate-batch")
async def translate_batch(
    files: list[UploadFile] = File(...),
    source_lang: str = Form("en"),
    target_lang: str = Form("ro"),
    provider: str = Form("auto"),
):
    """Traducere batch: pana la 5 fisiere, procesare secventiala, rezultate per-fisier."""
    if len(files) > 5:
        raise HTTPException(400, "Maximum 5 fisiere per batch")
    if not files:
        raise HTTPException(400, "Niciun fisier primit")

    results = []
    for f in files:
        fname = f.filename or "unknown"
        ext = Path(fname).suffix.lower()
        if ext not in (".pdf", ".docx", ".txt", ".md"):
            results.append({"filename": fname, "status": "error", "output_url": None, "error": f"Format nesuportat: {ext}"})
            continue

        tmp_path = await save_upload(f)
        try:
            if ext in (".txt", ".md"):
                with open(tmp_path, "r", encoding="utf-8", errors="ignore") as fh:
                    content = fh.read()
                result = await translate_with_chain(content, source_lang, target_lang, provider)
                translated = result["translated_text"]
                output_name = Path(fname).stem + f"_{target_lang.upper()}.txt"
                results.append({"filename": fname, "status": "ok", "output_name": output_name, "chars": len(content), "provider": result["provider"], "error": None})
            elif ext == ".docx":
                from docx import Document as DocxDocument
                doc = DocxDocument(tmp_path)
                total_chars = 0
                for para in doc.paragraphs:
                    if not para.text.strip():
                        continue
                    total_chars += len(para.text)
                    r = await translate_with_chain(para.text, source_lang, target_lang, provider)
                    if para.runs:
                        para.runs[0].text = r["translated_text"]
                        for run in para.runs[1:]:
                            run.text = ""
                    else:
                        para.text = r["translated_text"]
                output_name = Path(fname).stem + f"_{target_lang.upper()}.docx"
                results.append({"filename": fname, "status": "ok", "output_name": output_name, "chars": total_chars, "provider": provider, "error": None})
            elif ext == ".pdf":
                import fitz
                pdf_doc = fitz.open(tmp_path)
                total_chars = sum(len(page.get_text()) for page in pdf_doc)
                pdf_doc.close()
                output_name = Path(fname).stem + f"_{target_lang.upper()}.docx"
                results.append({"filename": fname, "status": "ok", "output_name": output_name, "chars": total_chars, "provider": provider, "error": None})
        except Exception as e:
            logger.error("Batch translate error for %s: %s", fname, e)
            results.append({"filename": fname, "status": "error", "output_name": None, "chars": 0, "provider": None, "error": str(e)[:200]})
        finally:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)

    await log_activity(
        action="translator.batch",
        summary=f"Batch traducere: {len(files)} fisiere, {sum(1 for r in results if r['status'] == 'ok')} ok",
        details={"files": [r["filename"] for r in results]},
    )

    return {"results": results, "total": len(results), "ok": sum(1 for r in results if r["status"] == "ok")}


# ---------------------------------------------------------------------------
# GET /api/translator/provider-stats — Provider latency statistics
# ---------------------------------------------------------------------------

@router.get("/provider-stats")
async def get_provider_stats():
    """R4-14: Statistici latenta per provider de traducere."""
    latency_store = get_provider_latencies()
    stats = []
    for prov_name in ["deepl", "azure", "google", "gemini", "openai"]:
        latencies = latency_store.get(prov_name, [])
        last = latencies[-1] if latencies else None
        avg = round(sum(latencies) / len(latencies), 1) if latencies else None
        stats.append({
            "provider": prov_name,
            "last_latency_ms": round(last, 1) if last is not None else None,
            "avg_latency_ms": avg,
            "measurements": len(latencies),
        })
    return stats


# ---------------------------------------------------------------------------
# POST /api/translator/quality-check — Translation quality assessment (AI)
# ---------------------------------------------------------------------------

@router.post("/quality-check")
async def quality_check(req: QualityCheckRequest):
    """Evalueaza calitatea unei traduceri folosind AI."""
    if not req.source_text.strip():
        raise HTTPException(400, "Textul sursa nu poate fi gol")
    if not req.translated_text.strip():
        raise HTTPException(400, "Textul tradus nu poate fi gol")

    from modules.ai.providers import ai_generate

    system_prompt = (
        "Esti un expert evaluator de traduceri profesionale. "
        "Raspunde STRICT in format JSON valid, fara markdown, fara backticks, fara explicatii in afara JSON-ului."
    )

    prompt = f"""Evalueaza calitatea urmatoarei traduceri din {req.source_lang.upper()} in {req.target_lang.upper()}.

Text original ({req.source_lang.upper()}):
{req.source_text}

Traducere ({req.target_lang.upper()}):
{req.translated_text}

Raspunde DOAR cu un JSON valid in acest format exact:
{{"score": <numar intreg 1-10>, "issues": ["problema 1", "problema 2"], "suggestions": ["sugestie 1", "sugestie 2"]}}

Criterii de evaluare:
- Acuratete (fidelitate fata de original)
- Fluenta (naturalete in limba tinta)
- Terminologie (termeni tehnici corecti)
- Gramatica si ortografie
- Stil si registru

Daca traducerea e perfecta, returneaza score 10 si liste goale.
Raspunde DOAR cu JSON-ul, nimic altceva."""

    try:
        result = await ai_generate(prompt, system_prompt=system_prompt)
        ai_text = result["text"].strip()
        provider = result["provider"]

        # Parse the AI response as JSON
        import json as json_mod

        # Clean up common AI response quirks
        if ai_text.startswith("```"):
            # Remove markdown code blocks
            ai_text = ai_text.strip("`").strip()
            if ai_text.startswith("json"):
                ai_text = ai_text[4:].strip()

        parsed = json_mod.loads(ai_text)

        score = max(1, min(10, int(parsed.get("score", 5))))
        issues = parsed.get("issues", [])
        suggestions = parsed.get("suggestions", [])

        # Ensure lists contain strings
        issues = [str(i) for i in issues if i] if isinstance(issues, list) else []
        suggestions = [str(s) for s in suggestions if s] if isinstance(suggestions, list) else []

    except json_mod.JSONDecodeError:
        # If AI didn't return valid JSON, return a basic assessment
        logger.warning("AI nu a returnat JSON valid pentru quality check: %s", ai_text[:200])
        score = 5
        issues = ["Evaluarea automata nu a putut fi procesata complet"]
        suggestions = ["Verificati manual traducerea"]
        provider = result.get("provider", "unknown") if "result" in dir() else "unknown"
    except RuntimeError as e:
        raise HTTPException(503, f"Niciun provider AI disponibil: {e}")
    except Exception as e:
        logger.error("Eroare la quality check: %s", e, exc_info=True)
        raise HTTPException(500, f"Eroare la evaluarea calitatii: {str(e)}")

    await log_activity(
        action="translator.quality_check",
        summary=f"Quality check: scor {score}/10 ({provider})",
        details={
            "score": score,
            "provider": provider,
            "source_lang": req.source_lang,
            "target_lang": req.target_lang,
            "source_chars": len(req.source_text),
            "issues_count": len(issues),
        },
    )

    return {
        "score": score,
        "issues": issues,
        "suggestions": suggestions,
        "provider": provider,
    }


# ---------------------------------------------------------------------------
# POST /api/translator/compare — Compare 2 providers side-by-side
# ---------------------------------------------------------------------------

@router.post("/compare")
async def compare_providers(req: CompareRequest):
    """Compara rezultatul traducerii de la 2 provideri diferiti, in paralel."""
    if not req.text.strip():
        raise HTTPException(400, "Textul nu poate fi gol")

    if req.provider_a == req.provider_b:
        raise HTTPException(400, "Alege doi provideri diferiti pentru comparatie")

    async def _translate_timed(provider_name: str) -> dict:
        """Translate with a specific provider and measure time."""
        start = time.perf_counter()
        try:
            result = await translate_with_chain(
                req.text, req.source_lang, req.target_lang, provider_name
            )
            elapsed = round(time.perf_counter() - start, 3)
            return {
                "provider": result["provider"],
                "translated_text": result["translated_text"],
                "time_seconds": elapsed,
                "error": None,
            }
        except Exception as e:
            elapsed = round(time.perf_counter() - start, 3)
            return {
                "provider": provider_name,
                "translated_text": None,
                "time_seconds": elapsed,
                "error": str(e),
            }

    result_a, result_b = await asyncio.gather(
        _translate_timed(req.provider_a),
        _translate_timed(req.provider_b),
    )

    await log_activity(
        action="translator.compare",
        summary=(
            f"Comparatie {req.provider_a} vs {req.provider_b} "
            f"({req.source_lang.upper()}->{req.target_lang.upper()}, {len(req.text)} car.)"
        ),
        details={
            "provider_a": req.provider_a,
            "provider_b": req.provider_b,
            "source_lang": req.source_lang,
            "target_lang": req.target_lang,
            "chars_count": len(req.text),
            "time_a": result_a["time_seconds"],
            "time_b": result_b["time_seconds"],
        },
    )

    return {
        "source_text": req.text,
        "source_lang": req.source_lang,
        "target_lang": req.target_lang,
        "chars_count": len(req.text),
        "provider_a": result_a,
        "provider_b": result_b,
    }
