"""
API endpoints for file conversion operations.

Endpoints:
  POST /api/converter/pdf-to-docx      — PDF → DOCX
  POST /api/converter/docx-to-pdf      — DOCX → PDF
  POST /api/converter/merge-pdfs       — multiple PDFs → single PDF
  POST /api/converter/split-pdf        — PDF → ZIP of individual pages
  POST /api/converter/compress-images  — reduce image file size
  POST /api/converter/resize-images    — resize image dimensions
  POST /api/converter/csv-to-json      — CSV → JSON
  POST /api/converter/excel-to-json    — XLSX → JSON
  POST /api/converter/create-zip       — multiple files → ZIP
  POST /api/converter/extract-text     — PDF/image → text (OCR)
"""

from __future__ import annotations

import asyncio
import csv
import io
import json
import logging
import os
import shutil
import uuid
import zipfile
from pathlib import Path
from tempfile import gettempdir

from fastapi import APIRouter, File, Form, HTTPException, UploadFile
from fastapi.responses import StreamingResponse

from app.core.activity_log import log_activity

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/converter", tags=["converter"])

_TMP = Path(gettempdir()) / "roland_converter"
_TMP.mkdir(exist_ok=True)

# ---------- File size limit ----------
_MAX_FILE_SIZE = 50 * 1024 * 1024  # 50 MB


def _check_file_size(content: bytes, filename: str | None = None):
    """Raise HTTP 413 if content exceeds _MAX_FILE_SIZE."""
    if len(content) > _MAX_FILE_SIZE:
        size_mb = len(content) / (1024 * 1024)
        raise HTTPException(
            413,
            f"Fisierul depaseste limita de 50MB. "
            f"Dimensiune: {size_mb:.1f}MB"
            + (f" ({filename})" if filename else ""),
        )


def _uid() -> str:
    return uuid.uuid4().hex[:8]


def _tmp(name: str) -> Path:
    return _TMP / f"{_uid()}_{name}"


def _cleanup(*paths: Path | str):
    for p in paths:
        try:
            p = Path(p)
            if p.is_file():
                p.unlink()
            elif p.is_dir():
                shutil.rmtree(p)
        except Exception:
            pass


def _stream(data: bytes, media_type: str, filename: str):
    return StreamingResponse(
        io.BytesIO(data),
        media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ---------- Android-safe validation helpers ----------

def _check_type(filename: str | None, content_type: str | None,
                extensions: tuple[str, ...], ct_keywords: tuple[str, ...]) -> bool:
    """Check file type by extension OR content-type (Android may strip extensions)."""
    name = (filename or "").lower()
    ct = (content_type or "").lower()
    if any(name.endswith(ext) for ext in extensions):
        return True
    if any(kw in ct for kw in ct_keywords):
        return True
    # application/octet-stream = generic binary, allow and let library decide
    if ct == "application/octet-stream":
        return True
    return False


def _safe_name(filename: str | None, default_ext: str) -> str:
    """Ensure filename has proper extension (Android may strip it)."""
    name = filename or f"file{default_ext}"
    if not Path(name).suffix:
        name += default_ext
    return name


_is_pdf = lambda fn, ct: _check_type(fn, ct, (".pdf",), ("pdf",))
_is_docx = lambda fn, ct: _check_type(fn, ct, (".docx", ".doc"), ("wordprocessing", "msword"))
_is_csv = lambda fn, ct: _check_type(fn, ct, (".csv",), ("csv",))
_is_excel = lambda fn, ct: _check_type(fn, ct, (".xlsx", ".xls"), ("spreadsheet", "excel"))
_is_image = lambda fn, ct: _check_type(fn, ct,
    (".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".webp", ".gif"),
    ("image/",))


# ---------- PDF → DOCX ----------

@router.post("/pdf-to-docx")
async def pdf_to_docx(file: UploadFile = File(...)):
    logger.info("PDF->DOCX request: filename=%r, content_type=%r", file.filename, file.content_type)

    if not _is_pdf(file.filename, file.content_type):
        raise HTTPException(400, f"Fisierul trebuie sa fie PDF. Primit: {file.filename!r} ({file.content_type})")

    content = await file.read()
    _check_file_size(content, file.filename)
    safe = _safe_name(file.filename, ".pdf")
    out_name = f"{Path(safe).stem}.docx"

    def _convert(data: bytes) -> bytes:
        from pdf2docx import Converter

        inp = _tmp(safe)
        out = _tmp(out_name)
        try:
            inp.write_bytes(data)
            cv = Converter(str(inp))
            cv.convert(str(out))
            cv.close()
            return out.read_bytes()
        finally:
            _cleanup(inp, out)

    try:
        result = await asyncio.to_thread(_convert, content)
    except Exception as e:
        logger.exception("PDF->DOCX failed")
        raise HTTPException(500, f"Conversie esuata: {e}")

    await log_activity(action="converter", summary=f"PDF -> DOCX: {file.filename}",
                       details={"type": "pdf-to-docx", "file": file.filename, "size": len(content)})

    return _stream(result,
                   "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                   out_name)


@router.post("/docx-to-pdf")
async def docx_to_pdf(file: UploadFile = File(...)):
    logger.info("DOCX->PDF request: filename=%r, content_type=%r, size_hint=%s",
                file.filename, file.content_type, file.size)

    if not _is_docx(file.filename, file.content_type):
        raise HTTPException(
            400,
            f"Fisierul trebuie sa fie DOCX/DOC. "
            f"Primit: filename={file.filename!r}, type={file.content_type!r}"
        )

    content = await file.read()
    _check_file_size(content, file.filename)
    logger.info("DOCX->PDF: read %d bytes", len(content))

    # Force .docx extension for temp file (Android may strip it)
    safe_name = file.filename or "document.docx"
    if not safe_name.lower().endswith((".docx", ".doc")):
        safe_name += ".docx"
    out_name = f"{Path(safe_name).stem}.pdf"

    def _convert_with_word(data: bytes) -> bytes:
        """Primary path: use COM automation via docx2pdf (requires MS Word)."""
        import pythoncom
        pythoncom.CoInitialize()
        try:
            from docx2pdf import convert

            inp = _tmp(safe_name)
            out = _tmp(out_name)
            try:
                inp.write_bytes(data)
                logger.info("DOCX->PDF: converting via Word %s -> %s", inp, out)
                convert(str(inp), str(out))
                if not out.exists():
                    raise FileNotFoundError(f"Output file not created: {out}")
                result = out.read_bytes()
                logger.info("DOCX->PDF: Word success, output %d bytes", len(result))
                return result
            finally:
                _cleanup(inp, out)
        finally:
            pythoncom.CoUninitialize()

    def _convert_fallback(data: bytes) -> bytes:
        """Fallback: extract text via python-docx, generate PDF via reportlab.

        Produces a basic text-only PDF (no complex formatting) but works
        without Microsoft Word installed.
        """
        from docx import Document
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table

        logger.info("DOCX->PDF: using reportlab fallback (Word unavailable)")

        doc = Document(io.BytesIO(data))

        # Build PDF
        buf = io.BytesIO()
        pdf = SimpleDocTemplate(
            buf,
            pagesize=A4,
            leftMargin=2 * cm,
            rightMargin=2 * cm,
            topMargin=2 * cm,
            bottomMargin=2 * cm,
        )

        styles = getSampleStyleSheet()
        # Create styles for headings with explicit fontName to avoid missing font errors
        heading_styles = {}
        for level in range(1, 4):
            size = {1: 16, 2: 14, 3: 12}[level]
            heading_styles[level] = ParagraphStyle(
                f"CustomHeading{level}",
                parent=styles["Normal"],
                fontSize=size,
                fontName="Helvetica-Bold",
                spaceAfter=6,
                spaceBefore=12,
            )
        normal_style = ParagraphStyle(
            "CustomNormal",
            parent=styles["Normal"],
            fontSize=10,
            fontName="Helvetica",
            spaceAfter=4,
        )

        story = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                story.append(Spacer(1, 6))
                continue
            # Escape XML special chars for reportlab
            text = (
                text.replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
            )
            style_name = (para.style.name or "").lower()
            if "heading 1" in style_name:
                story.append(Paragraph(text, heading_styles[1]))
            elif "heading 2" in style_name:
                story.append(Paragraph(text, heading_styles[2]))
            elif "heading 3" in style_name:
                story.append(Paragraph(text, heading_styles[3]))
            else:
                story.append(Paragraph(text, normal_style))

        # Extract tables
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    cell_text = (
                        cell_text.replace("&", "&amp;")
                        .replace("<", "&lt;")
                        .replace(">", "&gt;")
                    )
                    row_data.append(Paragraph(cell_text, normal_style))
                table_data.append(row_data)
            if table_data:
                t = Table(table_data)
                story.append(Spacer(1, 12))
                story.append(t)
                story.append(Spacer(1, 12))

        if not story:
            story.append(Paragraph("(Document gol)", normal_style))

        pdf.build(story)
        result = buf.getvalue()
        logger.info("DOCX->PDF: fallback success, output %d bytes", len(result))
        return result

    # Try Word COM first, fallback to reportlab
    use_fallback = False
    try:
        result = await asyncio.to_thread(_convert_with_word, content)
    except ImportError:
        logger.warning("DOCX->PDF: pythoncom/docx2pdf not available, using fallback")
        use_fallback = True
    except HTTPException:
        raise
    except Exception as e:
        logger.warning("DOCX->PDF: Word conversion failed (%s), trying fallback", e)
        use_fallback = True

    if use_fallback:
        try:
            result = await asyncio.to_thread(_convert_fallback, content)
        except Exception as e:
            logger.exception("DOCX->PDF fallback also failed")
            raise HTTPException(
                500,
                f"Conversie esuata. Word indisponibil, iar conversia de baza a esuat: {e}",
            )

    await log_activity(action="converter", summary=f"DOCX -> PDF: {file.filename}",
                       details={"type": "docx-to-pdf", "file": file.filename, "size": len(content)})

    return _stream(result, "application/pdf", out_name)


# ---------- Merge PDFs ----------

@router.post("/merge-pdfs")
async def merge_pdfs(files: list[UploadFile] = File(...)):
    logger.info("Merge PDFs: %d files", len(files))
    if len(files) < 2:
        raise HTTPException(400, "Trebuie minim 2 fisiere PDF")

    for f in files:
        if not _is_pdf(f.filename, f.content_type):
            raise HTTPException(400, f"Fisierul {f.filename!r} nu este PDF ({f.content_type})")

    file_data = []
    for f in files:
        data = await f.read()
        _check_file_size(data, f.filename)
        file_data.append((_safe_name(f.filename, ".pdf"), data))

    def _merge(items: list[tuple[str, bytes]]) -> bytes:
        import fitz

        merged = fitz.open()
        temps = []
        try:
            for name, data in items:
                tmp = _tmp(name)
                tmp.write_bytes(data)
                temps.append(tmp)
                doc = fitz.open(str(tmp))
                merged.insert_pdf(doc)
                doc.close()
            buf = merged.tobytes()
            merged.close()
            return buf
        finally:
            _cleanup(*temps)

    try:
        result = await asyncio.to_thread(_merge, file_data)
    except Exception as e:
        raise HTTPException(500, f"Merge esuat: {e}")

    names = [f.filename for f in files]
    await log_activity(action="converter", summary=f"Merge {len(files)} PDFs",
                       details={"type": "merge-pdfs", "files": names})

    return _stream(result, "application/pdf", "merged.pdf")


# ---------- Split PDF ----------

@router.post("/split-pdf")
async def split_pdf(
    file: UploadFile = File(...),
    pages: str = Form("all"),
):
    logger.info("Split PDF: filename=%r, content_type=%r", file.filename, file.content_type)
    if not _is_pdf(file.filename, file.content_type):
        raise HTTPException(400, f"Fisierul trebuie sa fie PDF. Primit: {file.filename!r} ({file.content_type})")

    content = await file.read()
    _check_file_size(content, file.filename)
    safe = _safe_name(file.filename, ".pdf")
    stem = Path(safe).stem

    def _split(data: bytes, page_spec: str) -> bytes:
        import fitz

        tmp = _tmp(safe)
        tmp.write_bytes(data)
        try:
            doc = fitz.open(str(tmp))
            total = len(doc)

            if page_spec.strip().lower() == "all":
                indices = list(range(total))
            else:
                indices = _parse_pages(page_spec, total)

            buf = io.BytesIO()
            with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
                for i in indices:
                    page_doc = fitz.open()
                    page_doc.insert_pdf(doc, from_page=i, to_page=i)
                    zf.writestr(f"{stem}_page_{i + 1}.pdf", page_doc.tobytes())
                    page_doc.close()
            doc.close()
            return buf.getvalue()
        finally:
            _cleanup(tmp)

    try:
        result = await asyncio.to_thread(_split, content, pages)
    except HTTPException:
        raise
    except ValueError as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        raise HTTPException(500, f"Split esuat: {e}")

    await log_activity(action="converter", summary=f"Split PDF: {file.filename} ({pages})",
                       details={"type": "split-pdf", "file": file.filename, "pages": pages})

    return _stream(result, "application/zip", f"{stem}_pages.zip")


def _parse_pages(spec: str, total: int) -> list[int]:
    """Parse page spec like '1,3,5-7' into 0-based indices.

    Raises ValueError with clear message on invalid input.
    """
    if not spec or not spec.strip():
        raise ValueError("Specificatia de pagini este goala")

    indices = []
    for part in spec.split(","):
        part = part.strip()
        if not part:
            continue
        if "-" in part:
            pieces = part.split("-", 1)
            start_s, end_s = pieces[0].strip(), pieces[1].strip()
            if not start_s or not end_s:
                raise ValueError(
                    f"Interval de pagini invalid: '{part}'. "
                    f"Format corect: '1-5'"
                )
            try:
                s = int(start_s)
                e = int(end_s)
            except ValueError:
                raise ValueError(
                    f"Valoare non-numerica in intervalul '{part}'. "
                    f"Folositi doar numere intregi, ex: '1-5'"
                )
            if s <= 0 or e <= 0:
                raise ValueError(
                    f"Numere de pagina negative sau zero in '{part}'. "
                    f"Paginile incep de la 1"
                )
            if s > e:
                raise ValueError(
                    f"Interval invalid: '{part}' — "
                    f"pagina de start ({s}) > pagina de sfarsit ({e})"
                )
            if s > total:
                raise ValueError(
                    f"Pagina {s} depaseste totalul de {total} pagini"
                )
            e = min(e, total)
            indices.extend(range(s - 1, e))
        else:
            try:
                p = int(part)
            except ValueError:
                raise ValueError(
                    f"Valoare non-numerica: '{part}'. "
                    f"Folositi doar numere intregi, ex: '1,3,5-7'"
                )
            if p <= 0:
                raise ValueError(
                    f"Numar de pagina invalid: {p}. Paginile incep de la 1"
                )
            if p > total:
                raise ValueError(
                    f"Pagina {p} depaseste totalul de {total} pagini"
                )
            indices.append(p - 1)

    if not indices:
        raise ValueError(
            f"Nicio pagina valida selectata din totalul de {total} pagini"
        )

    return sorted(set(indices))


# ---------- Compress Images ----------

_SUPPORTED_OUTPUT_FORMATS = {"jpeg", "webp", "png"}
_FORMAT_EXTENSIONS = {"jpeg": ".jpg", "webp": ".webp", "png": ".png"}
_FORMAT_MIMETYPES = {"jpeg": "image/jpeg", "webp": "image/webp", "png": "image/png"}


@router.post("/compress-images")
async def compress_images(
    files: list[UploadFile] = File(...),
    quality: int = Form(80),
    output_format: str = Form("jpeg"),
    report_only: bool = Form(False),
):
    quality = max(1, min(100, quality))
    output_format = output_format.lower().strip()
    if output_format not in _SUPPORTED_OUTPUT_FORMATS:
        raise HTTPException(
            400,
            f"Format nesuportat: {output_format!r}. "
            f"Formate acceptate: {', '.join(sorted(_SUPPORTED_OUTPUT_FORMATS))}",
        )

    file_data = []
    for f in files:
        data = await f.read()
        _check_file_size(data, f.filename)
        file_data.append((f.filename, data))
    single = len(files) == 1

    def _compress(
        items: list[tuple[str, bytes]], q: int, fmt: str
    ) -> tuple[bytes, str, list[dict]]:
        from PIL import Image

        ext = _FORMAT_EXTENSIONS[fmt]
        results = []
        report = []
        for name, data in items:
            original_size = len(data)
            img = Image.open(io.BytesIO(data))
            # WebP supports RGBA, JPEG does not
            if fmt == "jpeg" and img.mode in ("RGBA", "P", "LA"):
                img = img.convert("RGB")
            elif fmt == "png" and img.mode not in ("RGBA", "RGB", "L", "P"):
                img = img.convert("RGB")
            out = io.BytesIO()
            save_kwargs = {"format": fmt.upper(), "optimize": True}
            if fmt in ("jpeg", "webp"):
                save_kwargs["quality"] = q
            img.save(out, **save_kwargs)
            compressed_data = out.getvalue()
            compressed_size = len(compressed_data)
            reduction = (
                round((1 - compressed_size / original_size) * 100, 1)
                if original_size > 0
                else 0.0
            )
            out_name = f"{Path(name).stem}_compressed{ext}"
            results.append((out_name, compressed_data))
            report.append({
                "filename": name,
                "original_size": original_size,
                "compressed_size": compressed_size,
                "reduction_percent": reduction,
            })

        if len(results) == 1:
            return results[0][1], results[0][0], report

        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
            for name, data in results:
                zf.writestr(name, data)
        return buf.getvalue(), "compressed_images.zip", report

    try:
        result, out_name, report = await asyncio.to_thread(
            _compress, file_data, quality, output_format
        )
    except Exception as e:
        raise HTTPException(500, f"Compresie esuata: {e}")

    await log_activity(
        action="converter",
        summary=f"Compress {len(files)} images (q={quality}, fmt={output_format})",
        details={
            "type": "compress-images",
            "count": len(files),
            "quality": quality,
            "output_format": output_format,
            "report": report,
        },
    )

    # Report-only mode: return JSON stats without the file
    if report_only:
        total_original = sum(r["original_size"] for r in report)
        total_compressed = sum(r["compressed_size"] for r in report)
        total_reduction = (
            round((1 - total_compressed / total_original) * 100, 1)
            if total_original > 0
            else 0.0
        )
        return {
            "total_original_size": total_original,
            "total_compressed_size": total_compressed,
            "total_reduction_percent": total_reduction,
            "output_format": output_format,
            "quality": quality,
            "files": report,
        }

    # Normal mode: return file with compression headers
    media = _FORMAT_MIMETYPES.get(output_format, "image/jpeg") if single else "application/zip"
    response = _stream(result, media, out_name)
    # Add compression report as custom headers
    total_original = sum(r["original_size"] for r in report)
    total_compressed = sum(r["compressed_size"] for r in report)
    total_reduction = (
        round((1 - total_compressed / total_original) * 100, 1)
        if total_original > 0
        else 0.0
    )
    response.headers["X-Original-Size"] = str(total_original)
    response.headers["X-Compressed-Size"] = str(total_compressed)
    response.headers["X-Reduction-Percent"] = str(total_reduction)
    return response


# ---------- Resize Images ----------

@router.post("/resize-images")
async def resize_images(
    files: list[UploadFile] = File(...),
    width: int = Form(800),
    height: int = Form(0),
    keep_ratio: bool = Form(True),
):
    file_data = []
    for f in files:
        data = await f.read()
        _check_file_size(data, f.filename)
        file_data.append((f.filename, data))
    single = len(files) == 1

    def _resize(items: list[tuple[str, bytes]], w: int, h: int, ratio: bool) -> tuple[bytes, str]:
        from PIL import Image

        results = []
        for name, data in items:
            img = Image.open(io.BytesIO(data))
            orig_w, orig_h = img.size

            if ratio:
                if h <= 0:
                    scale = w / orig_w
                    new_w, new_h = w, int(orig_h * scale)
                elif w <= 0:
                    scale = h / orig_h
                    new_w, new_h = int(orig_w * scale), h
                else:
                    scale = min(w / orig_w, h / orig_h)
                    new_w, new_h = int(orig_w * scale), int(orig_h * scale)
            else:
                new_w = w if w > 0 else orig_w
                new_h = h if h > 0 else orig_h

            img = img.resize((new_w, new_h), Image.LANCZOS)
            out = io.BytesIO()
            fmt = img.format or ("PNG" if img.mode == "RGBA" else "JPEG")
            if fmt == "JPEG" and img.mode == "RGBA":
                img = img.convert("RGB")
            img.save(out, format=fmt, quality=90)
            ext = fmt.lower()
            out_name = f"{Path(name).stem}_{new_w}x{new_h}.{ext}"
            results.append((out_name, out.getvalue()))

        if len(results) == 1:
            return results[0][1], results[0][0]

        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
            for name, data in results:
                zf.writestr(name, data)
        return buf.getvalue(), "resized_images.zip"

    try:
        result, out_name = await asyncio.to_thread(_resize, file_data, width, height, keep_ratio)
    except Exception as e:
        raise HTTPException(500, f"Redimensionare esuata: {e}")

    media = "application/zip" if not single else "image/jpeg"
    await log_activity(action="converter", summary=f"Resize {len(files)} images ({width}x{height})",
                       details={"type": "resize-images", "count": len(files), "width": width, "height": height})

    return _stream(result, media, out_name)


# ---------- CSV → JSON ----------

@router.post("/csv-to-json")
async def csv_to_json(
    file: UploadFile = File(...),
    delimiter: str = Form(","),
):
    logger.info("CSV->JSON: filename=%r, content_type=%r", file.filename, file.content_type)
    if not _is_csv(file.filename, file.content_type):
        raise HTTPException(400, f"Fisierul trebuie sa fie CSV. Primit: {file.filename!r} ({file.content_type})")

    content = await file.read()
    _check_file_size(content, file.filename)
    safe = _safe_name(file.filename, ".csv")
    out_name = f"{Path(safe).stem}.json"

    def _convert(data: bytes, delim: str) -> bytes:
        text = data.decode("utf-8-sig")
        reader = csv.DictReader(io.StringIO(text), delimiter=delim)
        rows = list(reader)
        return json.dumps(rows, ensure_ascii=False, indent=2).encode("utf-8")

    try:
        result = await asyncio.to_thread(_convert, content, delimiter)
    except Exception as e:
        raise HTTPException(500, f"Conversie esuata: {e}")

    await log_activity(action="converter", summary=f"CSV -> JSON: {file.filename}",
                       details={"type": "csv-to-json", "file": file.filename})

    return _stream(result, "application/json", out_name)


# ---------- Excel → JSON ----------

@router.post("/excel-to-json")
async def excel_to_json(file: UploadFile = File(...)):
    logger.info("Excel->JSON: filename=%r, content_type=%r", file.filename, file.content_type)
    if not _is_excel(file.filename, file.content_type):
        raise HTTPException(400, f"Fisierul trebuie sa fie XLSX/XLS. Primit: {file.filename!r} ({file.content_type})")

    content = await file.read()
    _check_file_size(content, file.filename)
    safe = _safe_name(file.filename, ".xlsx")
    out_name = f"{Path(safe).stem}.json"

    def _convert(data: bytes) -> bytes:
        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        result = {}
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                result[sheet_name] = []
                continue
            headers = [str(h) if h is not None else f"col_{i}" for i, h in enumerate(rows[0])]
            data_rows = []
            for row in rows[1:]:
                obj = {}
                for i, val in enumerate(row):
                    key = headers[i] if i < len(headers) else f"col_{i}"
                    if val is not None:
                        obj[key] = val if not hasattr(val, "isoformat") else val.isoformat()
                    else:
                        obj[key] = None
                data_rows.append(obj)
            result[sheet_name] = data_rows
        wb.close()
        return json.dumps(result, ensure_ascii=False, indent=2).encode("utf-8")

    try:
        result = await asyncio.to_thread(_convert, content)
    except Exception as e:
        raise HTTPException(500, f"Conversie esuata: {e}")

    await log_activity(action="converter", summary=f"Excel -> JSON: {file.filename}",
                       details={"type": "excel-to-json", "file": file.filename})

    return _stream(result, "application/json", out_name)


# ---------- Create ZIP ----------

@router.post("/create-zip")
async def create_zip(files: list[UploadFile] = File(...)):
    if not files:
        raise HTTPException(400, "Trebuie cel putin un fisier")

    file_data = []
    for f in files:
        data = await f.read()
        _check_file_size(data, f.filename)
        file_data.append((f.filename, data))

    def _zip(items: list[tuple[str, bytes]]) -> bytes:
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
            for name, data in items:
                zf.writestr(name, data)
        return buf.getvalue()

    result = await asyncio.to_thread(_zip, file_data)

    await log_activity(action="converter", summary=f"ZIP: {len(files)} fisiere",
                       details={"type": "create-zip", "count": len(files),
                                "files": [f.filename for f in files]})

    return _stream(result, "application/zip", "archive.zip")


# ---------- Extract Text (OCR) ----------

@router.post("/extract-text")
async def extract_text(file: UploadFile = File(...)):
    logger.info("Extract text: filename=%r, content_type=%r", file.filename, file.content_type)
    content = await file.read()
    _check_file_size(content, file.filename)

    def _detect_and_extract(data: bytes, filename: str | None, ct: str | None) -> str:
        fn = (filename or "").lower()
        content_t = (ct or "").lower()
        if fn.endswith(".pdf") or "pdf" in content_t:
            return _extract_pdf_text(data)
        elif (fn.endswith((".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".webp", ".gif"))
              or content_t.startswith("image/")):
            return _extract_image_text(data)
        elif content_t == "application/octet-stream":
            # Try PDF first, then image
            try:
                return _extract_pdf_text(data)
            except Exception:
                return _extract_image_text(data)
        else:
            raise ValueError(f"Format nesuportat: {filename!r} ({ct}). Acceptate: PDF, PNG, JPG, TIFF, BMP, WEBP")

    try:
        text = await asyncio.to_thread(_detect_and_extract, content, file.filename, file.content_type)
    except ValueError as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        raise HTTPException(500, f"Extragere esuata: {e}")

    await log_activity(action="converter", summary=f"Extract text: {file.filename}",
                       details={"type": "extract-text", "file": file.filename, "chars": len(text)})

    safe = _safe_name(file.filename, ".txt")
    out_name = f"{Path(safe).stem}.txt"
    return _stream(text.encode("utf-8"), "text/plain; charset=utf-8", out_name)


def _extract_pdf_text(data: bytes) -> str:
    import fitz

    doc = fitz.open(stream=data, filetype="pdf")
    pages = []
    for i, page in enumerate(doc):
        text = page.get_text()
        if text.strip():
            pages.append(f"--- Pagina {i + 1} ---\n{text.strip()}")
    doc.close()

    if not pages:
        try:
            return _ocr_pdf(data)
        except Exception:
            return "(Nu s-a putut extrage text — PDF scanat fara OCR disponibil)"

    return "\n\n".join(pages)


def _extract_image_text(data: bytes) -> str:
    try:
        import pytesseract
        from PIL import Image

        img = Image.open(io.BytesIO(data))
        return pytesseract.image_to_string(img, lang="ron+eng")
    except ImportError:
        raise ValueError("pytesseract nu este instalat — OCR indisponibil")


def _ocr_pdf(data: bytes) -> str:
    import fitz

    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        raise ValueError("pytesseract nu este instalat")

    doc = fitz.open(stream=data, filetype="pdf")
    pages = []
    for i, page in enumerate(doc):
        pix = page.get_pixmap(dpi=300)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        text = pytesseract.image_to_string(img, lang="ron+eng")
        if text.strip():
            pages.append(f"--- Pagina {i + 1} ---\n{text.strip()}")
    doc.close()
    return "\n\n".join(pages) if pages else "(Nu s-a gasit text)"
