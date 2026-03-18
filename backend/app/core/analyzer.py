"""
Modul de extragere caracteristici din fișiere PDF și DOCX.

Analizează: pagini, cuvinte, imagini, tabele, diagrame, complexitate layout,
detecție documente scanate (OCR), densitate text.
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

import fitz  # PyMuPDF
import pdfplumber
from docx import Document as DocxDocument
from docx.opc.constants import RELATIONSHIP_TYPE as RT

logger = logging.getLogger(__name__)

# Tipuri de fișiere acceptate
SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


def extract_features(file_path: str | Path) -> dict[str, Any]:
    """
    Extrage caracteristici dintr-un fișier PDF sau DOCX.

    Returnează un dicționar cu:
        page_count, word_count, words_per_page, image_count, table_count,
        has_complex_tables, chart_count, has_diagrams, layout_complexity,
        is_scanned, text_density
    """
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext == ".pdf":
        return _extract_pdf_features(path)
    elif ext == ".docx":
        return _extract_docx_features(path)
    else:
        raise ValueError(
            f"Tip fișier nesuportat: {ext}. "
            f"Tipuri acceptate: {', '.join(SUPPORTED_EXTENSIONS)}"
        )


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def _extract_pdf_features(path: Path) -> dict[str, Any]:
    """Extrage caracteristici din PDF folosind PyMuPDF + pdfplumber."""
    features: dict[str, Any] = {}

    # --- PyMuPDF: text, imagini, pagini ---
    doc = fitz.open(str(path))
    page_count = len(doc)
    features["page_count"] = page_count

    total_words = 0
    total_images = 0
    total_chars = 0
    words_per_page_list: list[int] = []
    has_multi_column = False
    has_mixed_orientation = False
    prev_rotation = None

    for page in doc:
        text = page.get_text("text")
        words = text.split()
        page_word_count = len(words)
        total_words += page_word_count
        total_chars += len(text)
        words_per_page_list.append(page_word_count)

        # Imagini pe pagină
        image_list = page.get_images(full=True)
        total_images += len(image_list)

        # Detectare coloane multiple: verificăm blocurile de text
        blocks = page.get_text("blocks")
        text_blocks = [b for b in blocks if b[6] == 0]  # tip 0 = text
        if len(text_blocks) > 1:
            x_positions = sorted(set(round(b[0], -1) for b in text_blocks))
            if len(x_positions) >= 2:
                # Dacă blocurile încep la x-uri semnificativ diferite
                if x_positions[-1] - x_positions[0] > page.rect.width * 0.3:
                    has_multi_column = True

        # Detectare orientare mixtă
        rotation = page.rotation
        if prev_rotation is not None and rotation != prev_rotation:
            has_mixed_orientation = True
        prev_rotation = rotation

    doc.close()

    features["word_count"] = total_words
    features["words_per_page"] = (
        round(total_words / page_count, 1) if page_count > 0 else 0
    )
    features["image_count"] = total_images

    # --- pdfplumber: tabele (cu timeout per pagină) ---
    table_count = 0
    complex_table_count = 0
    chart_count = 0

    try:
        with pdfplumber.open(str(path)) as pdf:
            # Limităm la primele 5 pagini — find_tables() e foarte lent pe PDF-uri mari
            max_pages = min(len(pdf.pages), 5)
            for page_idx in range(max_pages):
                page = pdf.pages[page_idx]
                try:
                    tables = page.find_tables()
                    page_table_count = len(tables)
                    table_count += page_table_count

                    for table in tables:
                        extracted = table.extract()
                        if extracted:
                            num_rows = len(extracted)
                            num_cols = len(extracted[0]) if extracted[0] else 0
                            if num_cols > 5 or num_rows > 10:
                                complex_table_count += 1

                    if page.images:
                        for img in page.images:
                            img_width = abs(img.get("x1", 0) - img.get("x0", 0))
                            img_height = abs(img.get("top", 0) - img.get("bottom", 0))
                            page_area = page.width * page.height
                            if page_area > 0:
                                img_area_ratio = (img_width * img_height) / page_area
                                if img_area_ratio > 0.20:
                                    chart_count += 1
                except Exception as page_exc:
                    logger.warning("pdfplumber eroare pe pagina %d: %s", page_idx, page_exc)
                    continue
    except Exception as exc:
        logger.warning("pdfplumber nu a putut analiza tabelele: %s", exc)

    features["table_count"] = table_count
    features["has_complex_tables"] = complex_table_count > 0
    features["chart_count"] = chart_count

    # Detectare diagrame/scheme: bazat pe nr mare de imagini + context tehnic
    features["has_diagrams"] = total_images > page_count * 0.5 or chart_count > 2

    # --- Layout complexity (1-5) ---
    complexity = 1
    if has_multi_column:
        complexity += 1
    if has_mixed_orientation:
        complexity += 1
    if table_count > 3:
        complexity += 1
    if total_images > page_count:
        complexity += 1
    features["layout_complexity"] = min(complexity, 5)

    # --- OCR detection ---
    avg_words_per_page = features["words_per_page"]
    features["is_scanned"] = avg_words_per_page < 10

    # --- Text density ---
    # Caractere per pagină (normalizat la o pagină A4 tipică ~2000 chars)
    chars_per_page = total_chars / page_count if page_count > 0 else 0
    features["text_density"] = round(min(chars_per_page / 2000.0, 2.0), 2)

    return features


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def _extract_docx_features(path: Path) -> dict[str, Any]:
    """Extrage caracteristici din DOCX folosind python-docx."""
    doc = DocxDocument(str(path))
    features: dict[str, Any] = {}

    # --- Text și cuvinte ---
    total_words = 0
    paragraph_texts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraph_texts.append(text)
            total_words += len(text.split())

    # Estimare pagini: ~300 cuvinte/pagină pentru documente tehnice
    estimated_pages = max(1, round(total_words / 300))
    features["page_count"] = estimated_pages
    features["word_count"] = total_words
    features["words_per_page"] = (
        round(total_words / estimated_pages, 1) if estimated_pages > 0 else 0
    )

    # --- Imagini ---
    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_count += 1
    features["image_count"] = image_count

    # --- Tabele ---
    table_count = len(doc.tables)
    complex_table_count = 0
    for table in doc.tables:
        num_rows = len(table.rows)
        num_cols = len(table.columns)
        if num_cols > 5 or num_rows > 10:
            complex_table_count += 1

    features["table_count"] = table_count
    features["has_complex_tables"] = complex_table_count > 0

    # --- Grafice și diagrame ---
    # În DOCX, graficele sunt obiecte inline — numărăm imaginile mari
    features["chart_count"] = 0  # Dificil de detectat precis în DOCX
    features["has_diagrams"] = image_count > estimated_pages * 0.5

    # --- Layout complexity ---
    complexity = 1
    # Verificare secțiuni cu coloane
    for section in doc.sections:
        # python-docx nu expune direct nr coloane, dar putem verifica lățimea
        pass
    if table_count > 3:
        complexity += 1
    if image_count > estimated_pages:
        complexity += 1
    if complex_table_count > 0:
        complexity += 1
    features["layout_complexity"] = min(complexity, 5)

    # --- OCR / scanat ---
    features["is_scanned"] = False  # DOCX-urile nu sunt de obicei scanate

    # --- Text density ---
    total_chars = sum(len(t) for t in paragraph_texts)
    chars_per_page = total_chars / estimated_pages if estimated_pages > 0 else 0
    features["text_density"] = round(min(chars_per_page / 2000.0, 2.0), 2)

    return features


def get_feature_vector(features: dict[str, Any]) -> list[float]:
    """
    Convertește dicționarul de caracteristici într-un vector numeric
    pentru algoritmii de machine learning.

    Ordinea: page_count, word_count, words_per_page, image_count,
             table_count, has_complex_tables, chart_count, has_diagrams,
             layout_complexity, is_scanned, text_density
    """
    return [
        float(features.get("page_count", 0)),
        float(features.get("word_count", 0)),
        float(features.get("words_per_page", 0)),
        float(features.get("image_count", 0)),
        float(features.get("table_count", 0)),
        float(features.get("has_complex_tables", False)),
        float(features.get("chart_count", 0)),
        float(features.get("has_diagrams", False)),
        float(features.get("layout_complexity", 1)),
        float(features.get("is_scanned", False)),
        float(features.get("text_density", 0)),
    ]
